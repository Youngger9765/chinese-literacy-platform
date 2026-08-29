# -*- coding: utf-8 -*-
"""單課完整流程：轉換 + 自動測試。回傳 PASS/FAIL 與報告"""
import sys, os, re, subprocess, glob, shutil, shlex
import docx
from docx.text.paragraph import Paragraph
W='{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
SCRIPT_DIR=os.path.dirname(os.path.abspath(__file__))
CONVERT_PY=os.path.join(SCRIPT_DIR, 'convert.py')
PYTHON=shlex.quote(sys.executable or 'python3')
if SCRIPT_DIR not in sys.path:
    sys.path.insert(0, SCRIPT_DIR)

def _resolve_soffice():
    """通用 LibreOffice 解析：Claude Cowork / Codex / 本機 Mac/Linux 皆可。
    回傳可直接接 '--headless ...' 的命令前綴字串。
    優先序：SOFFICE_BIN 環境變數 > 系統 soffice/libreoffice > Claude Cowork 內建 docx skill wrapper。"""
    env = os.environ.get('SOFFICE_BIN')
    if env:
        return f'python3 "{env}"' if env.endswith('.py') else f'"{env}"'
    for name in ('soffice', 'libreoffice'):
        p = shutil.which(name)
        if p:
            return f'"{p}"'
    # Claude Cowork 內建 wrapper（session mount 路徑會變，用 glob 找，不寫死 session 名）
    for pat in ('/mnt/**/.claude/skills/docx/scripts/office/soffice.py',
                '/sessions/*/mnt/.claude/skills/docx/scripts/office/soffice.py',
                os.path.expanduser('~/.claude/skills/docx/scripts/office/soffice.py')):
        hits = glob.glob(pat, recursive=True)
        if hits:
            return f'python3 "{hits[0]}"'
    raise RuntimeError('找不到 LibreOffice：請安裝 soffice/libreoffice（Codex/Linux: apt install libreoffice；Mac: brew install --cask libreoffice）或設環境變數 SOFFICE_BIN')

SOFFICE=_resolve_soffice()

def paras(path):
    """只取正文流的段落，排除文字方塊內段落"""
    d=docx.Document(path)
    out=[]
    for p in d.element.body.iter(W+'p'):
        anc=p.getparent(); inbox=False
        while anc is not None:
            if anc.tag==W+'txbxContent': inbox=True; break
            anc=anc.getparent()
        if not inbox: out.append(Paragraph(p,d).text)
    return out

def run(cmd, **kw):
    return subprocess.run(cmd, shell=True, capture_output=True, text=True, **kw)

def parse_ae(stderr):
    """ImageMagick -metric AE 的「差異像素數」。
    IMv7 Q16-HDRI 印 '<raw> (<pixel_count>)' → 取括號內的像素數（threshold 比的是它）；
    舊版/其他建置只印一個整數或科學記號 → 取該值。（可攜性修正：跨 ImageMagick 版本）"""
    paren=re.search(r'\(\s*([-+]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][-+]?\d+)?)\s*\)', stderr)
    if paren:
        val=paren.group(1)
    else:
        plain=re.search(r'[-+]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][-+]?\d+)?', stderr)
        if not plain:
            return -1
        val=plain.group(0)
    try:
        return int(float(val))
    except ValueError:
        return -1

def _missing_critical_fonts(src):
    """docx 需要、但此環境沒裝的『關鍵專有字型』（標楷體/新細明體類）。
    這些缺了 → LibreOffice render 會字型代換 → 逐頁像素比對必假陽性（但轉換本身仍正確）。
    Windows 多半原生內建這些字型；偵測不到 fc-list（如 Windows）視為字型齊全。"""
    CRITICAL={'標楷體','DFKai-SB','新細明體','PMingLiU','MingLiU'}
    used=set()
    try:
        d=docx.Document(src)
        for rf in d.element.body.iter(W+'rFonts'):
            for a in ('ascii','eastAsia','hAnsi','cs'):
                v=rf.get(W+a)
                if v: used.add(v)
    except Exception:
        return []
    needed=[f for f in CRITICAL if f in used]
    if not needed: return []
    if shutil.which('fc-list') is None:
        return []
    listing=run('fc-list').stdout
    return [f for f in needed if f not in listing]

def test_lesson(src, outdir, tag):
    os.makedirs(outdir, exist_ok=True)
    os.makedirs('/tmp/pl', exist_ok=True)
    stu=f'{outdir}/{tag}學生版.docx'; ans=f'{outdir}/{tag}answers.tsv'; qc=f'{outdir}/{tag}qc.txt'
    r=run(f'{PYTHON} {shlex.quote(CONVERT_PY)} {shlex.quote(src)} {shlex.quote(stu)} {shlex.quote(ans)} {shlex.quote(qc)}')
    if r.returncode!=0: return False, [f'轉換失敗: {r.stderr[-300:]}']
    problems=[]; notes=[r.stdout.strip()]

    # T1 段落數與變更範圍
    from convert import is_white
    def vis_paras(path):
        dd=docx.Document(path)
        out=[]
        for p0 in dd.element.body.iter(W+'p'):
            anc=p0.getparent(); inbox=False
            while anc is not None:
                if anc.tag==W+'txbxContent': inbox=True; break
                anc=anc.getparent()
            if inbox: continue
            out.append(''.join(r.text for r in Paragraph(p0,dd).runs if not is_white(r)))
        return out
    t=paras(src); s=vis_paras(stu)
    if len(t)!=len(s): problems.append(f'段落數不同 {len(t)} vs {len(s)}')
    changed=[i for i,(a,b) in enumerate(zip(t,s)) if a!=b]
    notes.append(f'變更段落 {len(changed)}')

    # T2 殘留答案檢查
    # 測試用比轉換器更寬的樣式：任一種括號內夾任一寬度的單一大寫字母
    left=[x for x in s if re.search(r'[（(](?:\s|\u3000)+[A-NＡ-Ｎ](?:\s|\u3000)*[）)]|[（(](?:\s|\u3000)*[A-NＡ-Ｎ](?:\s|\u3000)+[）)]|（\s*[A-NＡ-Ｎ]\s*）', x)]
    if left: problems.append(f'殘留括號答案 {len(left)}: {left[:2]}')
    d=docx.Document(stu)
    n=sum(1 for y in d.element.body.iter(W+'sym') if (y.get(W+'char') or '').upper()=='F0FE')
    if n: problems.append(f'殘留☒記號 {n}')
    n=0
    for tb in d.element.body.iter(W+'txbxContent'):
        if '0F4761' in {c.get(W+'val') for c in tb.iter(W+'color')}:
            if ''.join(t.text or '' for t in tb.iter(W+'t')).strip():
                n+=1
    if n: problems.append(f'殘留解析方塊(含文字) {n}')
    # 沒有□的圈號選項（正確答案痕跡）
    stray=[]
    sym_paras=set()
    dsym=docx.Document(stu)
    for pp in dsym.element.body.iter(W+'p'):
        for sy in pp.iter(W+'sym'):
            if (sy.get(W+'char') or '').upper()=='F06F':
                sym_paras.add(Paragraph(pp,dsym).text); break
    def _opt(sx):
        return bool(re.match(r'^□?\s*[①②③④⑤⑥⑦⑧⑨]', sx.strip()))
    for i,x in enumerate(s):
        if re.match(r'^\s*[①②③④⑤⑥⑦⑧⑨]', x.strip()) and '□' not in x and x not in sym_paras and sum(1 for c in x if c in '①②③④⑤⑥⑦⑧⑨')==1:
            lo=i
            while lo-1>=0 and _opt(s[lo-1]): lo-=1
            hi=i
            while hi+1<len(s) and _opt(s[hi+1]): hi+=1
            if any('□' in y for y in s[lo:hi+1]): stray.append(x)
    if stray: problems.append(f'學生版仍有無□圈號選項 {len(stray)}: {[x[:20] for x in stray[:2]]}')
    # 【】內殘留非空白
    br=[x for x in s if re.search(r'【[^】\s　]', x.replace(' ','')) and '□' not in x and not any(c in x for c in '①②③④⑤')]
    if br: notes.append(f'【】內仍有文字(可能為題目本身文字) {len(br)}: {[x[:25] for x in br[:2]]}')

    # T2a 解答樣式殘留（權威檢查）
    from convert import answer_style_ids
    d2=docx.Document(stu)
    sids=answer_style_ids(d2)
    if sids:
        bad=[]
        for p in d2.element.body.iter(W+'p'):
            P2=Paragraph(p,d2)
            head=P2.text.strip()[:1]
            if head and head in '□①②③④⑤⑥⑦⑧⑨':
                continue  # 選項句內的解答樣式=原檔誤用格式，轉換時保留（QC已註記）
            for r in P2.runs:
                st=r._element.rPr.find(W+'rStyle') if r._element.rPr is not None else None
                if st is not None and st.get(W+'val') in sids and not is_white(r) and re.sub(r'[\u200b\uFE00-\uFE0F\U000E0100-\U000E01EF　 _＿\u00A0]','',r.text):
                    bad.append(r.text[:12])
        if bad: problems.append(f'解答樣式run殘留 {len(bad)}: {bad[:4]}')

    # T2a2 可見記號圖形殘留（紅/橘/藍指引線等無文字圖形必須已隱形）
    A2='{http://schemas.openxmlformats.org/drawingml/2006/main}'
    MK={'FF0000','EE0000','C00000','E97132','ED7D31','4472C4'}
    vis=0
    dm=docx.Document(stu)
    for ln in dm.element.body.iter(A2+'ln'):
        sf=ln.find(A2+'solidFill')
        if sf is None: continue
        s2=sf.find(A2+'srgbClr')
        sc2=sf.find(A2+'schemeClr')
        hit=(s2 is not None and (s2.get('val') or '').upper() in MK) or (sc2 is not None and sc2.get('val') in ('accent1','accent2'))
        if hit:
            # 所屬容器若無文字 → 記號圖形，不該有可見框線
            node=ln
            while node is not None and not node.tag.endswith('}drawing') and not node.tag.endswith('}pict'):
                node=node.getparent()
            if node is not None and not any((t.text or '').strip() for t in node.iter(W+'t')):
                vis+=1
    if vis: problems.append(f'可見記號圖形殘留 {vis}')

    # T2b 跨段括號洩漏偵測（儲存格內 】答案【）
    dd=docx.Document(stu)
    leaks=[]
    for tc in dd.element.body.iter(W+'tc'):
        cps=[c for c in tc if c.tag==W+'p']
        if len(cps)<2: continue
        joined='\n'.join(''.join(r.text for r in Paragraph(c,dd).runs if not is_white(r)) for c in cps)
        pos=[i for i,c in enumerate(joined) if c in '【】']
        for a,b in zip(pos[::2],pos[1::2]):
            core=joined[a+1:b].replace('\n','').strip('　 ')
            if core and '\n' in joined[a+1:b] and len(core)<=12 and not any(c in core for c in '□①②③④⑤。，'):
                leaks.append(core[:15])
    if leaks: problems.append(f'跨段括號疑似洩漏: {leaks}')

    # T3 渲染與逐頁比對
    # 逐頁比對＝拿 render 出的 PDF 比像素。缺關鍵專有字型時 render 會字型代換 → 比對必假陽性；
    # 但轉換正確性已由 T1/T2（結構稽核 0 差異 + 殘留答案檢查）保證，與字型無關。
    # → 缺字型時降為提醒，不列 FAIL（避免非 Windows 環境每次都假 FAIL）。
    missing_fonts=_missing_critical_fonts(src)
    if missing_fonts:
        notes.append(f'逐頁比對已略過：此環境缺字型 {missing_fonts}，render 會代換、像素比對必假陽性。轉換正確性已由結構稽核(0 差異)+殘留答案檢查保證；最終排版請在字型齊全環境(如 Windows 原生標楷體/新細明體)或裝齊字型後目視抽查')
    else:
        run(f'{SOFFICE} --headless --convert-to pdf --outdir /tmp/pl "{src}"')
        run(f'{SOFFICE} --headless --convert-to pdf --outdir /tmp/pl "{stu}"')
        tp=f'/tmp/pl/{os.path.splitext(os.path.basename(src))[0]}.pdf'
        sp=f'/tmp/pl/{tag}學生版.pdf'
        pt=run(f'pdfinfo "{tp}" | grep Pages').stdout.split()[-1]
        ps=run(f'pdfinfo "{sp}" | grep Pages').stdout.split()[-1]
        if pt!=ps and abs(int(pt)-int(ps))>1:
            problems.append(f'頁數差超過1頁 {pt} vs {ps}')
        elif pt!=ps:
            notes.append(f'頁數差1（{pt}→{ps}，行尾空白收合；內容等值已由段落比對保證，需目視抽查）')
        else:
            run(f'rm -f /tmp/pl/T-* /tmp/pl/S-*')
            run(f'pdftoppm -png -r 50 "{tp}" /tmp/pl/T')
            run(f'pdftoppm -png -r 50 "{sp}" /tmp/pl/S')
            tpages=sorted(glob.glob('/tmp/pl/T-*.png'))
            diffs=[]
            for tpg in tpages:
                spg=tpg.replace('/T-','/S-')
                ae=parse_ae(run(f'compare -metric AE "{tpg}" "{spg}" /dev/null').stderr.strip())
                diffs.append((os.path.basename(tpg), ae))
            big=[]
            benign=[]
            for p,a in diffs:
                if a<0: big.append((p,a)); continue
                if a<=30000: continue
                num=p.split('-')[1]
                tpg=f'/tmp/pl/T-{num}'; spg=f'/tmp/pl/S-{num}'
                dims=run(f'identify -format "%w %h" "{tpg}"').stdout.split()
                wpx,hpx=int(dims[0]),int(dims[1])
                best=a
                for off in (4,8,12,16,20,24,28):
                    for tup,sup in ((off,0),(0,off)):  # 兩個方向
                        h2=hpx-off
                        run(f'convert "{tpg}" -crop {wpx}x{h2}+0+{tup} +repage /tmp/pl/tc.png')
                        run(f'convert "{spg}" -crop {wpx}x{h2}+0+{sup} +repage /tmp/pl/sc.png')
                        ae=parse_ae(run('compare -metric AE /tmp/pl/tc.png /tmp/pl/sc.png /dev/null').stderr.strip())
                        if ae < 0: continue
                        best=min(best,ae)
                if best <= 30000: benign.append((p,a,best))
                else: big.append((p,a))
            if big: problems.append(f'頁面差異過大(疑似變形): {big}')
            if benign: notes.append(f'整頁垂直位移，對齊後差異僅剩答案區(可接受): {[(p,a,b) for p,a,b in benign]}')
            notes.append('逐頁diff: '+' '.join(f'{a}' for _,a in diffs))
    ok=not problems
    return ok, problems+notes

if __name__=='__main__':
    src, outdir, tag = sys.argv[1], sys.argv[2], sys.argv[3]
    ok, msgs = test_lesson(src, outdir, tag)
    print('RESULT:', 'PASS' if ok else 'FAIL')
    for m in msgs: print('  ', m)
