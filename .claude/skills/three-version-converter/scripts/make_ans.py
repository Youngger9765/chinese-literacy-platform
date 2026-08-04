# -*- coding: utf-8 -*-
import sys, re, docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build(lesson, tsv, out):
    rows = [l.rstrip('\n').split('\t') for l in open(tsv) if l.strip()]
    d = docx.Document()
    st = d.styles['Normal']; st.font.name='Microsoft JhengHei'; st.font.size=Pt(11)
    st.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia','Microsoft JhengHei')
    h = d.add_paragraph(); r = h.add_run(lesson + ' 簡答'); r.bold=True; r.font.size=Pt(16)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 分組
    vocab, cloze, mc, mid = [], [], [], []
    for k,c,a in rows:
        if k=='教師解析(移除)':
            continue
        if k=='填空(底線)' and re.match(r'^\(\d+\)', c.strip()) and '：' in c:
            vocab.append(a)
        elif k in ('語詞應用',) or (k=='選擇題' and re.match(r'^\(\d+\)', c.strip())):
            num = re.match(r'^\((\d+)\)', c.strip())
            cloze.append((num.group(1) if num else '?', a))
        elif k=='選擇題' and c.startswith('第'):
            mc.append((c, a))
        else:
            mid.append((k,c,a))

    def head(t):
        p=d.add_paragraph(); r=p.add_run('◎'+t); r.bold=True; r.font.size=Pt(12)
    if vocab:
        head('語詞我最棒')
        d.add_paragraph('　'.join(f'({i+1}) {w}' for i,w in enumerate(vocab)))
    if cloze:
        head('語詞應用')
        d.add_paragraph('　'.join(f'({n}) {a}' for n,a in cloze))
    if mid:
        head('文章重點表／閱讀聚光燈')
        for k,c,a in mid:
            c1 = re.sub(r'\s+',' ',c)[:38]
            d.add_paragraph(f'・{a}　（{c1}）')
    if mc:
        head('閱讀理解')
        d.add_paragraph('　'.join(f'{c.replace("第","").replace("題","")}. {a}' for c,a in mc))
    d.save(out)
    print(out, 'ok')

build(sys.argv[1], sys.argv[2], sys.argv[3])
