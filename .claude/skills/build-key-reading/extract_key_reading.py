#!/usr/bin/env python3
"""重點朗讀段抽取器（實驗版）— 從教師版 DOCX 抽「念順順」重點段.

演算法（2026-07-20 對 G6-L22 驗證過）:
  1. 找「朗讀表格」= 某 table 的 cell 含累計字數欄（一串遞增數字，末值 ~300-350）
  2. 課文欄 = 該 row 內文字最長的 cell；累計字數欄 = 含遞增數字的 cell
  3. ☞ 起點 = 課文欄中「含 w:drawing/w:pict」的段落（手指頭圖形）
  4. 長度 benchmark = max(累計字數)（= 不含標點字數）
  5. 從起點取字到不含標點字數 ≈ benchmark，再 snap 到下一個句尾（。」！？）
     → 顯示段落（不斷句）；benchmark 保留原值供 CPM 計分
"""
import sys, re
import docx

PUNCT = '，。、！？：；「」『』（）〈〉《》…—　 \n\t.'
SENT_END = '。」！？'


def strip_punct(s: str) -> str:
    return re.sub(r'[，。、！？：；「」『』（）〈〉《》…—\s.]', '', s)


def find_reading_table(doc):
    """回傳 (table, count_col_idx, text_col_idx) 或 None."""
    for t in doc.tables:
        for row in t.rows:
            cells = row.cells
            # 累計字數欄 = 一堆遞增數字、末值在 150-600 之間
            for ci, cell in enumerate(cells):
                nums = [int(x) for x in re.findall(r'\d+', cell.text)]
                if len(nums) >= 4 and nums == sorted(nums) and 150 <= nums[-1] <= 600:
                    # 課文欄 = 同 row 內文字最長的 cell（排除自己）
                    text_ci = max(
                        (i for i in range(len(cells)) if i != ci),
                        key=lambda i: len(strip_punct(cells[i].text)),
                    )
                    return t, row, ci, text_ci, nums
    return None


def find_start_paragraph(text_cell):
    """☞ 起點 = 含圖形的段落 index；找不到回 None."""
    for pi, p in enumerate(text_cell.paragraphs):
        xml = p._p.xml
        if 'w:drawing' in xml or 'w:pict' in xml or 'w:object' in xml:
            return pi
    return None


def extract(docx_path):
    doc = docx.Document(docx_path)
    found = find_reading_table(doc)
    if not found:
        return {'ok': False, 'reason': 'no reading table (no cumulative-count column)'}
    t, row, count_ci, text_ci, nums = found
    benchmark = nums[-1]
    text_cell = row.cells[text_ci]
    start_pi = find_start_paragraph(text_cell)
    full = ''.join(p.text for p in text_cell.paragraphs)
    if start_pi is None:
        return {'ok': False, 'reason': 'no ☞ drawing in text column', 'benchmark': benchmark}
    start_char = sum(len(p.text) for p in text_cell.paragraphs[:start_pi])
    body = full[start_char:]

    # 取到不含標點字數 ≈ benchmark，再 snap 到句尾
    acc = 0
    cut = len(body)
    for i, ch in enumerate(body):
        if ch not in PUNCT:
            acc += 1
        if acc >= benchmark:
            cut = i + 1
            break
    # snap forward 到下一個句尾
    snap = cut
    while snap < len(body) and body[snap - 1] not in SENT_END:
        snap += 1
    passage = body[:snap].strip()

    return {
        'ok': True,
        'benchmark': benchmark,
        'start_text': body[:15],
        'passage': passage,
        'passage_chars_no_punct': len(strip_punct(passage)),
        'ends_clean': passage[-1] in SENT_END if passage else False,
        'raw_at_benchmark': body[:cut].strip()[-12:],
    }


if __name__ == '__main__':
    import json
    r = extract(sys.argv[1])
    print(json.dumps(r, ensure_ascii=False, indent=2))
