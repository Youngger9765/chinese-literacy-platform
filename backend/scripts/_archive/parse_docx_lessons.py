#!/usr/bin/env python3
"""
parse_docx_lessons.py
解析陳教授課程 docx → lesson YAML

用法:
    python3 backend/scripts/parse_docx_lessons.py \
        --source private/curriculum-source/2026-05-01/自學教材兩個單元 \
        --output backend/data/lessons \
        --images backend/data/images

支援:
    - G6 單元 A（4 課）: 摘要策略 問題-解決-結果結構
    - G7 單元 B（3 課）: 圖文整合閱讀策略
    - 圖片 extract（G7-L29 為主）→ backend/data/images/{lesson_code}/

Schema 沿用 L01.yml，新增欄位:
    - reading_strategy_type: "summary_psr" | "graphic_text_integration"
    - strategy_exercises: 圖文整合練習題（G7 新欄位）
    - images: [{filename, caption, related_paragraph}]（G7 新欄位）
    - story_structure: 問題/解決/結果表格（G6 保留）
    - discussion_questions: 進階小試身手題（G6）
    - self_check_items: 自我檢核清單
    - video_links: 影片連結
"""

import argparse
import os
import re
import sys
from pathlib import Path

import docx
import yaml

# ---------- helpers ----------

def clean_text(text: str) -> str:
    """移除多餘空白和全形空格"""
    return re.sub(r"\s+", " ", text.replace("　", " ")).strip()


def extract_answer_letter(text: str) -> str:
    """從 (  F  ) 或 （ C ） 或 (F) 等格式中抽出答案字母"""
    m = re.search(r"[（(]\s*([A-Za-zＡ-Ｚ])\s*[）)]", text)
    if m:
        return m.group(1).upper().translate(str.maketrans("ＡＢＣＤ", "ABCD"))
    return ""


def strip_answer_prefix(text: str) -> str:
    """移除題目前方的 （ C ）1. → 1."""
    return re.sub(r"^[（(]\s*[A-Za-zＡ-ＺＤ１２３４]\s*[）)]\s*", "", text).strip()


def parse_vocab_from_paragraphs(paragraphs: list[docx.text.paragraph.Paragraph]) -> list[dict]:
    """從樣式 '教材：(1)＿：填空題目' 的段落中解析生字"""
    vocab = []
    pattern = re.compile(r"^\((\d+)\)\s+(.+?)[:：]\s+(.+)$")
    for p in paragraphs:
        style_name = p.style.name if p.style else ""
        if "填空題目" not in style_name:
            continue
        text = clean_text(p.text)
        m = pattern.match(text)
        if m:
            word = clean_text(m.group(2))
            definition = clean_text(m.group(3))
            vocab.append({"word": word, "definition": definition})
    return vocab


def parse_vocab_bank(table: docx.table.Table) -> dict[str, str]:
    """解析題庫表格 Table (2×5 or similar) → {A: 詞語, B: 詞語, ...}"""
    bank = {}
    for row in table.rows:
        for cell in row.cells:
            text = clean_text(cell.text)
            m = re.match(r"^([A-J])\s*[.．]\s*(.+)$", text)
            if m:
                bank[m.group(1)] = clean_text(m.group(2))
    return bank


def parse_fill_in_blank(paragraphs: list, vocab_bank: dict[str, str]) -> list[dict]:
    """解析填空題段落 → [{sentence, answer, answer_word}]"""
    items = []
    pattern = re.compile(r"^\(\d+\)\s+(.+)$")
    for p in paragraphs:
        style = p.style.name if p.style else ""
        if style not in ("Normal", "", "List Paragraph"):
            continue
        text = clean_text(p.text)
        if not pattern.match(text):
            continue
        # 答案在括號中
        ans = extract_answer_letter(text)
        if not ans:
            continue
        # 移除題號 (1)
        sentence = re.sub(r"^\(\d+\)\s+", "", text).strip()
        items.append({
            "sentence": sentence,
            "answer": ans,
            "answer_word": vocab_bank.get(ans, ""),
        })
    return items


def parse_multiple_choice(paragraphs: list, start_idx: int = 0) -> list[dict]:
    """解析選擇題 → [{question, options:[A,B,C,D], answer}]"""
    questions = []
    i = start_idx
    while i < len(paragraphs):
        p = paragraphs[i]
        text = clean_text(p.text)
        ans = extract_answer_letter(text)
        # 題目行：以 （X）N. 開頭
        q_match = re.match(r"^[（(]\s*[A-Da-dＡ-ＤＢＣ１２３４]\s*[）)]\s*\d+[.．]", text)
        if q_match and ans:
            question_text = strip_answer_prefix(text)
            options = []
            i += 1
            while i < len(paragraphs):
                opt_p = paragraphs[i]
                opt_text = clean_text(opt_p.text)
                # 選項行：以 A. B. C. D. 開頭（可能在同一行多個）
                if re.match(r"^[A-Da-d][.．]", opt_text):
                    # 可能一行多個 "A.xxx  B.xxx"
                    parts = re.split(r"\s{2,}(?=[A-Da-d][.．])", opt_text)
                    for part in parts:
                        part = part.strip()
                        if re.match(r"^[A-Da-d][.．]", part):
                            options.append(re.sub(r"^[A-Da-d][.．]\s*", "", part))
                    i += 1
                elif re.match(r"^[（(]\s*[A-Da-dＡ-ＤＢＣ]\s*[）)]\s*\d+", opt_text):
                    # 下一題了
                    break
                else:
                    i += 1
                    break
            questions.append({
                "question": question_text,
                "options": options,
                "answer": ans,
            })
            continue
        i += 1
    return questions


def parse_story_structure_table(table: docx.table.Table) -> list[dict]:
    """解析問題-解決-結果表格"""
    rows_data = []
    for row in table.rows:
        cells = [clean_text(c.text) for c in row.cells]
        # 去重（merged cells 會重複）
        unique = []
        for c in cells:
            if not unique or c != unique[-1]:
                unique.append(c)
        if any(unique):
            rows_data.append(unique)
    return rows_data


def parse_reading_benchmark(table: docx.table.Table) -> dict:
    """解析朗讀速率標準表格"""
    if len(table.rows) < 2:
        return {}
    thresholds = [clean_text(c.text) for c in table.rows[0].cells]
    feedbacks = [clean_text(c.text) for c in table.rows[1].cells]
    levels = []
    for t, f in zip(thresholds, feedbacks):
        if t and f:
            levels.append({"threshold": t, "feedback": f})
    return {"levels": levels} if levels else {}


def parse_video_links(table: docx.table.Table) -> list[dict]:
    """解析影片連結表"""
    links = []
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci == 0:
                continue
            text = clean_text(cell.text)
            if not text:
                continue
            lines = [l.strip() for l in text.split("\n") if l.strip()]
            if not lines:
                continue
            title = lines[0]
            url = ""
            for line in lines:
                if "http" in line:
                    url = line
                    break
            if title:
                links.append({"title": title, "url": url})
    return links


def extract_images_from_doc(doc: docx.Document, output_dir: Path, lesson_code: str) -> list[dict]:
    """Extract embedded images from docx to output_dir"""
    output_dir.mkdir(parents=True, exist_ok=True)
    image_list = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img_data = rel.target_part.blob
            ext = rel.target_part.content_type.split("/")[-1]
            if ext == "jpeg":
                ext = "jpg"
            filename = f"{lesson_code}-{len(image_list) + 1:02d}.{ext}"
            out_path = output_dir / filename
            with open(out_path, "wb") as f:
                f.write(img_data)
            image_list.append({
                "filename": f"images/{lesson_code}/{filename}",
                "caption": "",
            })
    return image_list


def get_story_text_from_table0(table: docx.table.Table) -> tuple[str, list[str], str, str]:
    """從 Table 0 (課文主表) 抽出 story_text、paragraphs、authors"""
    title = ""
    story_text = ""
    authors = ""
    if not table.rows:
        return title, [], authors, story_text

    # Row 0 Col 0 = 課號+課名
    title_cell = clean_text(table.rows[0].cells[0].text)
    m = re.match(r"第(\d+)課\s+(.+)", title_cell)
    if m:
        title = clean_text(m.group(2))

    # Row 0 Col 2 = 作者
    if len(table.rows[0].cells) >= 3:
        authors = clean_text(table.rows[0].cells[2].text)

    # Row 2 Col 1 = 課文
    if len(table.rows) >= 3 and len(table.rows[2].cells) >= 2:
        story_text = clean_text(table.rows[2].cells[1].text)

    paras = [p.strip() for p in story_text.split("  ") if p.strip()]
    if len(paras) <= 1:
        # 段落以 \n\n 分割
        paras = [p.strip() for p in story_text.replace("。\n", "。\n\n").split("\n\n") if p.strip()]

    return title, paras, authors, story_text


# ---------- G6 parser (摘要策略) ----------

def parse_g6_lesson(doc_path: Path, lesson_code: str) -> dict:
    doc = docx.Document(str(doc_path))
    paras = doc.paragraphs
    tables = doc.tables

    result = {
        "lesson_code": lesson_code,
        "grade": 6,
        "grade_code": lesson_code,
        "reading_strategy": "摘要策略－問題·解決·結果結構",
        "reading_strategy_type": "summary_psr",
        "genre": "說明文",
        "source_file": doc_path.name,
        "flags": [],
    }

    # --- Table 0: 課文 ---
    if tables:
        title, story_paras, authors, story_text = get_story_text_from_table0(tables[0])
        result["title"] = title
        result["authors"] = authors
        result["story_text"] = story_text
        result["paragraphs"] = story_paras
        result["paragraph_count"] = len(story_paras)
        result["char_count"] = len(story_text.replace(" ", ""))

    # --- 生字 (from paragraphs with style '教材：(1)＿：填空題目') ---
    vocab = parse_vocab_from_paragraphs(paras)
    result["vocabulary"] = vocab
    result["vocabulary_count"] = len(vocab)

    # --- 找計時表 (Table 1 or 2 depending on lesson) ---
    # --- 找朗讀速率表 ---
    for t in tables:
        cells_flat = [clean_text(c.text) for row in t.rows for c in row.cells]
        if any("還要多加練習" in c for c in cells_flat):
            result["reading_benchmark"] = parse_reading_benchmark(t)
            break

    # --- 影片連結 ---
    for t in tables:
        if t.rows and "影片連結" in clean_text(t.rows[0].cells[0].text):
            result["video_links"] = parse_video_links(t)
            break

    # --- 詞語題庫表 (2×5 or 2×N with A.詞語 format) ---
    vocab_bank = {}
    for t in tables:
        for row in t.rows:
            for cell in row.cells:
                m = re.match(r"^[A-J][.．]\s*.+", clean_text(cell.text))
                if m:
                    vocab_bank = parse_vocab_bank(t)
                    break
            if vocab_bank:
                break

    result["vocab_bank"] = vocab_bank

    # --- 填空題 ---
    fill_items = parse_fill_in_blank(paras, vocab_bank)
    result["fill_in_blank"] = fill_items

    # --- 選擇題 ---
    # 選擇題從大約 paragraph 70+ 開始
    mc = parse_multiple_choice(paras, start_idx=60)
    result["multiple_choice"] = mc
    result["multiple_choice_count"] = len(mc)

    # --- 問題解決結果表格 ---
    for t in tables:
        header_cells = [clean_text(c.text) for c in t.rows[0].cells] if t.rows else []
        if any("問題" in c or "解決" in c or "結果" in c for c in header_cells):
            result["story_structure_table"] = parse_story_structure_table(t)
            break
        # L22 has it in Row 1
        if len(t.rows) > 1:
            cells = [clean_text(c.text) for c in t.rows[1].cells]
            if any("問題" in c for c in cells):
                result["story_structure_table"] = parse_story_structure_table(t)
                break

    # --- 自我檢核 ---
    self_check = []
    in_check = False
    for p in paras:
        text = clean_text(p.text)
        if "自我檢核" in text:
            in_check = True
            continue
        if in_check and re.match(r"^[□☐]?\s*\d+\.", text):
            item = re.sub(r"^[□☐]?\s*\d+\.\s*", "", text).strip()
            if item:
                self_check.append(item)
        elif in_check and text.startswith("◎"):
            in_check = False
    result["self_check_items"] = self_check

    # --- 小試身手 discussion questions ---
    discussion = []
    in_discussion = False
    for p in paras:
        text = clean_text(p.text)
        if "小試身手" in text or "進階挑戰" in text:
            in_discussion = True
            continue
        if in_discussion and text.startswith("◎"):
            in_discussion = False
        if in_discussion and text:
            discussion.append(text)
    result["discussion_content"] = discussion

    return result


# ---------- G7 parser (圖文整合) ----------

def parse_g7_lesson(doc_path: Path, lesson_code: str, images_base: Path) -> dict:
    doc = docx.Document(str(doc_path))
    paras = doc.paragraphs
    tables = doc.tables

    result = {
        "lesson_code": lesson_code,
        "grade": 7,
        "grade_code": lesson_code,
        "reading_strategy": "圖文整合閱讀策略",
        "reading_strategy_type": "graphic_text_integration",
        "genre": "說明文",
        "source_file": doc_path.name,
        "flags": [],
    }

    # --- Table 1: 課文主表（G7 沒有 Table 0 vocab，Table 0 是說明框）---
    course_table_idx = 1  # G7 的課文在 tables[1]
    if len(tables) > course_table_idx:
        t = tables[course_table_idx]
        title, story_paras, authors, story_text = get_story_text_from_table0(t)
        result["title"] = title
        result["authors"] = authors
        result["story_text"] = story_text
        result["paragraphs"] = story_paras
        result["paragraph_count"] = len(story_paras)
        result["char_count"] = len(story_text.replace(" ", ""))

    # --- 影片連結 ---
    for t in tables:
        if t.rows and "影片連結" in clean_text(t.rows[0].cells[0].text):
            result["video_links"] = parse_video_links(t)
            break

    # --- 選擇題 ---
    mc = parse_multiple_choice(paras, start_idx=60)
    result["multiple_choice"] = mc
    result["multiple_choice_count"] = len(mc)

    # --- 策略練習題 (步驟❶❷❸❹ + 練習一/二/三/四) ---
    strategy_exercises = []
    current_exercise = None
    for p in paras:
        text = clean_text(p.text)
        # 新練習單元
        m_ex = re.match(r"^(練習[一二三四五])[:：]?\s*(.*)$", text)
        if m_ex:
            if current_exercise:
                strategy_exercises.append(current_exercise)
            current_exercise = {
                "exercise": m_ex.group(1),
                "description": clean_text(m_ex.group(2)),
                "steps": [],
            }
            continue
        # 步驟
        m_step = re.match(r"^(步驟[❶❷❸❹①②③④])[:：]?\s*(.+)$", text)
        if m_step and current_exercise is not None:
            current_exercise["steps"].append({
                "step": m_step.group(1),
                "description": m_step.group(2),
            })
            continue
    if current_exercise:
        strategy_exercises.append(current_exercise)
    result["strategy_exercises"] = strategy_exercises

    # --- 圖文架構摘要 (最後一個 table 通常是) ---
    for t in tables:
        if len(t.rows) >= 4 and len(t.columns) >= 2:
            cells = [clean_text(c.text) for c in t.rows[0].cells]
            # 如果第一行是課名重複（merged header）
            if len(t.rows) >= 3:
                second_row = [clean_text(c.text) for c in t.rows[1].cells]
                if any(kw in " ".join(second_row) for kw in ["研究問題", "假說", "問題", "解決"]):
                    result["graphic_text_summary_table"] = parse_story_structure_table(t)
                    break

    # --- 自我檢核 ---
    self_check = []
    in_check = False
    for p in paras:
        text = clean_text(p.text)
        if "自我檢核" in text:
            in_check = True
            continue
        if in_check and re.match(r"^[□☐]?\s*\d+\.", text):
            item = re.sub(r"^[□☐]?\s*\d+\.\s*", "", text).strip()
            if item:
                self_check.append(item)
        elif in_check and text.startswith("◎"):
            in_check = False
    result["self_check_items"] = self_check

    # --- 說明框 (Table 0 or 2) ---
    strategy_notes = []
    for t in [tables[0]] if tables else []:
        text = clean_text(t.rows[0].cells[0].text) if t.rows else ""
        if text:
            strategy_notes.append(text)
    result["strategy_intro_notes"] = strategy_notes

    # --- 圖片 extract ---
    img_output_dir = images_base / lesson_code
    images = extract_images_from_doc(doc, img_output_dir, lesson_code)
    result["images"] = images
    result["image_count"] = len(images)

    return result


# ---------- YAML 序列化 ----------

def represent_str(dumper, data):
    """多行字串用 | 格式"""
    if "\n" in data or len(data) > 80:
        return dumper.represent_scalar("tag:yaml.org,2002:str", data, style="|")
    return dumper.represent_scalar("tag:yaml.org,2002:str", data)


yaml.add_representer(str, represent_str)


def save_yaml(data: dict, path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        yaml.dump(data, f, allow_unicode=True, default_flow_style=False, sort_keys=False, width=120)


# ---------- main ----------

DOCX_MAP = {
    "G6-L22": "G6-L22小兵立大功：雞鳴狗盜的故事（摘要策略-問題.解決.結果結構）.docx",
    "G6-L23": "G6-L23老鷹紅豆的故事（摘要策略-問題.解決結構找重點）.docx",
    "G6-L24": "G6-L24白鯨救援（摘要策略-問題.解決結構找重點）.docx",
    "G6-L25": "G6-L25全世界第一張股票的誕生（摘要策略-問題.解決結構找重點）.docx",
    "G7-L28": "G7-L28看不見的兇手：以實驗破解肉湯腐敗之謎（圖文整合閱讀策略）.docx",
    "G7-L29": "G7-L29四張圖看地球暖化（圖文整合閱讀策略）.docx",
    "G7-L30": "G7-L30都是八哥為什麼命運不一樣（圖文表整合閱讀策略）.docx",
}


def main():
    parser = argparse.ArgumentParser(description="解析陳教授課程 docx → lesson YAML")
    parser.add_argument("--source", default="private/curriculum-source/2026-05-01/自學教材兩個單元")
    parser.add_argument("--output", default="backend/data/lessons")
    parser.add_argument("--images", default="backend/data/images")
    parser.add_argument("--lesson", default="all", help="指定課程代碼，如 G6-L22，或 all")
    args = parser.parse_args()

    source_dir = Path(args.source)
    output_dir = Path(args.output)
    images_dir = Path(args.images)

    if not source_dir.exists():
        print(f"ERROR: source dir not found: {source_dir}", file=sys.stderr)
        sys.exit(1)

    lessons_to_process = DOCX_MAP if args.lesson == "all" else {
        args.lesson: DOCX_MAP[args.lesson]
    }

    audit_rows = []

    for lesson_code, filename in lessons_to_process.items():
        doc_path = source_dir / filename
        if not doc_path.exists():
            print(f"WARN: docx not found: {doc_path}", file=sys.stderr)
            continue

        print(f"Parsing {lesson_code} ...")
        grade = int(lesson_code[1])

        if grade == 6:
            data = parse_g6_lesson(doc_path, lesson_code)
        elif grade == 7:
            data = parse_g7_lesson(doc_path, lesson_code, images_dir)
        else:
            print(f"WARN: unsupported grade {grade}", file=sys.stderr)
            continue

        out_path = output_dir / f"{lesson_code}.yml"
        save_yaml(data, out_path)
        print(f"  -> {out_path}")

        # Audit row
        audit_rows.append({
            "lesson": lesson_code,
            "title": data.get("title", ""),
            "vocab_count": data.get("vocabulary_count", 0),
            "fill_in_blank_count": len(data.get("fill_in_blank", [])),
            "mc_count": data.get("multiple_choice_count", 0),
            "has_images": data.get("image_count", 0) > 0,
            "image_count": data.get("image_count", 0),
            "has_story_structure": "story_structure_table" in data or "graphic_text_summary_table" in data,
            "has_video_links": len(data.get("video_links", [])) > 0,
            "self_check_count": len(data.get("self_check_items", [])),
            "strategy_exercises": len(data.get("strategy_exercises", [])),
        })

    # Print audit report
    print("\n" + "=" * 80)
    print("AUDIT REPORT")
    print("=" * 80)
    print(f"{'課程':<10} {'課名':<25} {'生字':<5} {'填空':<5} {'選擇':<5} {'圖片':<6} {'結構表':<6} {'影片':<5} {'自檢':<5}")
    print("-" * 80)
    for row in audit_rows:
        print(
            f"{row['lesson']:<10} "
            f"{row['title'][:22]:<25} "
            f"{row['vocab_count']:<5} "
            f"{row['fill_in_blank_count']:<5} "
            f"{row['mc_count']:<5} "
            f"{row['image_count']:<6} "
            f"{'Y' if row['has_story_structure'] else 'N':<6} "
            f"{'Y' if row['has_video_links'] else 'N':<5} "
            f"{row['self_check_count']:<5}"
        )
    print("=" * 80)
    print(f"Total: {len(audit_rows)} lessons processed.")


if __name__ == "__main__":
    main()
