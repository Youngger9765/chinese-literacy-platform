#!/usr/bin/env python3
"""
parse_docx_lessons_bulk.py
批量解析 158 課教師版 docx → YAML metadata

用法:
    python3 backend/scripts/parse_docx_lessons_bulk.py \
        --source "private/curriculum-source/2026-05-01/1.L1-158新版完成學習單1150415" \
        --output backend/data/lessons/_parsed_2026-05-01 \
        --images backend/data/lessons/_parsed_2026-05-01/images

功能:
    - 遞迴掃描 1-1教師版(L1~122)/ 所有子目錄 + 第三階段(L123開始~L157)差學生版/
    - 每個 docx → YAML + 圖片 extract
    - 錯誤逐檔 catch → 記入 parse-report.md，繼續其餘檔案
    - 冪等：已存在的 YAML 跳過（除非 --force）
    - 無 LLM 呼叫，純 docx 解析

改善項目（vs PR #1361 版）:
    - 修正 merged-cell Table 0 標題空白問題（G7-L29 等）
    - 新增 image metadata：paragraph_index, size, hash（供 #1341 chrome 過濾）
    - G6 詞語從 inline-style 段落補充解析
    - 影片連結從純文字段落 fallback 抽取 URL
    - 通用解析器（不限 G6/G7，支援 G4~G9 + 文言文）
"""

import argparse
import hashlib
import io
import os
import re
import sys
import traceback
from datetime import datetime
from pathlib import Path
from typing import Optional

import docx
import yaml

# ---------- constants ----------

TEACHER_DIR_NAME = "1-1教師版(L1~122)"
STAGE3_DIR_NAME = "第三階段(L123開始~L157)差學生版"

# Grade code mapping from directory name
GRADE_MAP = {
    "4年級": "G4",
    "5年級": "G5",
    "6年級": "G6",
    "7年級": "G7",
    "8年級": "G8",
    "9年級": "G9",
    "文言文": "文",
}

# ---------- text helpers ----------

def clean_text(text: str) -> str:
    """移除多餘空白和全形空格"""
    if not text:
        return ""
    return re.sub(r"\s+", " ", text.replace("　", " ")).strip()


# ---------- caption (圖N / 表N 圖說) helpers (#2218) ----------
#
# 圖文整合課文的「圖N …／表N …」圖說/表說行，在 docx 課文表格 cell 內是獨立一行。
# 舊版 split_story_paragraphs 會把整行當成課文段落留在 paragraphs[]，導致圖說在閱讀頁
# 被當成課文內文 render（#2218），而 images[].caption 反而是空的。
#
# 判別規則必須與前端 utils/paragraphMarkers.ts 一致：
#   圖說行 = ^(圖|表)[一二三…十]<空白>非空白...
#   行首標記後「緊接空白」才是圖說；緊接 CJK（例「表一比較了…」「圖三進一步說明…」）
#   是正常課文句子，**不可**剝離。

CN_NUMERALS: dict[str, int] = {
    "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
    "六": 6, "七": 7, "八": 8, "九": 9, "十": 10,
}
# 行首「圖N␠」/「表N␠」(N 為中文數字)；標記後必須是空白 + 非空白內容才算圖說行。
CAPTION_ROW_RE = re.compile(r"^(圖|表)([一二三四五六七八九十])[\s　]+(\S.*)$")


def detect_caption_row(text: str) -> Optional[dict]:
    """偵測單行是否為「圖N …／表N …」圖說/表說行。

    Returns {kind: '圖'|'表', number: int, label: '圖一', caption: str} or None.
    與前端 detectImageMarker/detectTableMarker 同規則（標記後需緊接空白）。
    """
    if not text:
        return None
    m = CAPTION_ROW_RE.match(text.strip())
    if not m:
        return None
    kind, cn, rest = m.group(1), m.group(2), m.group(3).strip()
    num = CN_NUMERALS.get(cn)
    if num is None or not rest:
        return None
    return {"kind": kind, "number": num, "label": f"{kind}{cn}", "caption": rest}


def split_caption_rows(paragraphs: list[str]) -> tuple[list[str], list[dict]]:
    """從段落清單分離「圖N／表N」圖說行 (#2218)。

    Returns (body_paragraphs, caption_rows)，caption_rows 為 detect_caption_row 的 dict。
    body 段落仍保留句中以 inline 方式提及的「如圖一」「從表二可知」等（不是行首圖說）。
    """
    body: list[str] = []
    captions: list[dict] = []
    for p in paragraphs:
        cap = detect_caption_row(p)
        if cap is not None:
            captions.append(cap)
        else:
            body.append(p)
    return body, captions


def assign_captions_to_images(images: list[dict], caption_rows: list[dict]) -> None:
    """把「圖N」圖說寫進對應 image 的 caption / figure_label (#2218)。

    對應策略（與前端 ComprehensionLayout.buildFigureIndex 一致）：
      1. 若 image 已有 figure_label（如「圖一」），以該標記的數字配對。
      2. 否則用陣列序位 (idx+1) 當 fallback figure number。
    只填空的 caption，不覆寫既有值。表說（表N）不寫進 images（表格另有 tables[].title）。
    """
    if not images:
        return
    image_caps = {c["number"]: c["caption"] for c in caption_rows if c["kind"] == "圖"}
    if not image_caps:
        return
    for idx, img in enumerate(images):
        label = img.get("figure_label") or ""
        num = None
        m = re.search(r"圖\s*([一二三四五六七八九十]|\d+)", label)
        if m:
            tok = m.group(1)
            num = CN_NUMERALS.get(tok) or (int(tok) if tok.isdigit() else None)
        if num is None:
            num = idx + 1  # positional fallback
        cap = image_caps.get(num)
        if cap and not img.get("caption"):
            img["caption"] = cap
            if not img.get("figure_label"):
                cn = next((k for k, v in CN_NUMERALS.items() if v == num), None)
                if cn:
                    img["figure_label"] = f"圖{cn}"


def extract_lesson_code_from_filename(filename: str) -> Optional[str]:
    """
    從檔名抽出 lesson code，如:
      G4-L1贏得喝采的輸家.docx → G4-L1
      G9-L15~16多文本-....docx → G9-L15-16
      文-L1假新聞？....docx    → 文-L1
      Ｇ8-L11....docx         → G8-L11  (全型 Ｇ)
    """
    stem = Path(filename).stem
    # 全型英文字母轉半型
    stem_normalized = stem.translate(str.maketrans("ＧＡＢ０１２３４５６７８９", "GAB0123456789"))

    # 模式: (G4|G5|...|文)-L(\d+[~\-]?\d*)
    m = re.match(r"^([Gg]?\d|文).*?[-－](L\d+[~～\-]?\d*)", stem_normalized)
    if m:
        prefix = m.group(1).upper()
        if prefix.isdigit():
            # Shouldn't happen but guard
            prefix = "G" + prefix
        lesson_part = m.group(2).replace("~", "-").replace("～", "-")
        return f"{prefix}-{lesson_part}"

    # fallback: match GN-LN directly
    m2 = re.match(r"^([G文]\d*-L[\d\-~～]+)", stem_normalized)
    if m2:
        code = m2.group(1).replace("~", "-").replace("～", "-")
        return code

    return None


def get_title_from_table0(table: docx.table.Table) -> tuple[str, str, str]:
    """
    從 Table 0 抽取 (title, authors, story_text)

    固定格式:
      Row 0, Col 0 = 合併儲存格，包含 "第N課  課名"
      Row 0, Col 2 = 作者（有些課沒有）
      Row 2, Col 1 = 課文（有些是 Col 0 或合併到 Col 2）

    改善: 遍歷所有 row[0] 的儲存格找課名，而非固定 [0][0]
    """
    title = ""
    authors = ""
    story_text = ""

    if not table.rows:
        return title, authors, story_text

    # 找課名 — 遍歷 Row 0 所有 unique cells
    seen_cells = set()
    for cell in table.rows[0].cells:
        txt = clean_text(cell.text)
        if txt in seen_cells:
            continue
        seen_cells.add(txt)

        if not title:
            # 模式1: "第N課  課名" or "第N 課  課名" (with optional space before 課)
            # Also handles "第？課" and "第? 課" (pending course number)
            m = re.match(r"第([?？\d]+)\s*課\s+(.+)", txt)
            if m:
                title = clean_text(m.group(2))
                # 清理課名中的長空格
                title = re.sub(r"\s{2,}", " ", title)
                continue

            # 模式2: 課名在課號後，但無明確 "第" 字（少數課）
            m2 = re.match(r"(\d+)\s+(.+)", txt)
            if m2 and len(m2.group(2)) > 3:
                title = clean_text(m2.group(2))

        # 找作者 — 包含 "：" 且有文字
        if not authors and "：" in txt and len(txt) < 60:
            # "課文/學習單：張明珠" 或 "課文：曾世杰  學習單：巫怡蓉"
            if re.search(r"[課文作者學習單]", txt):
                authors = txt

    # 找課文 — Row 2, Col 1（標準）
    # NOTE: 保留換行符號（\n），讓 split_story_paragraphs 能按段分割。
    #       只清除每行內部的多餘空格，不壓縮跨行的 \n。
    if len(table.rows) >= 3:
        row2_cells = table.rows[2].cells
        # 跳過第一欄（通常是段落標號）
        for col_idx in range(1, len(row2_cells)):
            cell_txt = row2_cells[col_idx].text.strip()
            if len(cell_txt) > 30:  # 課文通常較長
                # 保留換行：只清理每一行的首尾空白和行內多餘空格
                lines = [re.sub(r"[ \t　]+", " ", line).strip()
                         for line in cell_txt.split("\n")]
                story_text = "\n".join(line for line in lines if line)
                break
        # 如果 col 1 是空白，試 col 0
        if not story_text and row2_cells:
            cell_txt0 = row2_cells[0].text.strip()
            lines = [re.sub(r"[ \t　]+", " ", line).strip()
                     for line in cell_txt0.split("\n")]
            story_text = "\n".join(line for line in lines if line)

    return title, authors, story_text


def split_story_paragraphs(story_text: str) -> list[str]:
    """拆分課文段落

    優先策略（按順序，第一個產出 >=2 段即停）:
    1. 換行分段（docx table cell 的自然換行）
    2. 雙空格分段（部分格式用空格取代換行）
    3. 句號+換行分段
    """
    if not story_text:
        return []
    # 策略1: 換行分段（docx table cell 原生換行，最可靠）
    paras = [p.strip() for p in story_text.split("\n") if p.strip()]
    if len(paras) >= 2:
        return paras
    # 策略2: 雙空格分段
    paras = [p.strip() for p in story_text.split("  ") if p.strip()]
    if len(paras) >= 2:
        return paras
    # 策略3: 句號+換行分段
    paras = [p.strip() for p in story_text.replace("。\n", "。\n\n").split("\n\n") if p.strip()]
    if len(paras) >= 2:
        return paras
    # 最後: 整段當單一段落
    return [story_text.strip()] if story_text.strip() else []


# ---------- field extractors ----------

def extract_fill_in_blank(tables: list, story_table_idx: int = 0) -> list[dict]:
    """
    從學習單 tables 抽取填空題（【answer】格式）。

    跳過 story_table_idx 所指的課文主表（Table 0 或 1），
    掃描其餘 tables 的所有 cell，以 regex 抽出 【...】 markers。

    Returns list of:
      {
        "id": int (1-based),
        "answer": str,
        "context_before": str (前 50 字，供 OMO AI 比對位置),
        "context_after": str (後 50 字),
        "context_paragraph_idx": int (該 blank 在第幾個 cell 段),
      }
    """
    BLANK_RE = re.compile(r"【([^】]*)】")
    results = []
    blank_id = 1

    for ti, t in enumerate(tables):
        if ti == story_table_idx:
            continue  # 跳過課文主表
        for row in t.rows:
            # Deduplicate merged cells by tracking seen cell ids
            seen_cell_ids = set()
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cell_ids:
                    continue
                seen_cell_ids.add(cell_id)

                cell_text = cell.text
                for m in BLANK_RE.finditer(cell_text):
                    answer = clean_text(m.group(1))
                    start = m.start()
                    end = m.end()
                    before = clean_text(cell_text[max(0, start - 50):start])
                    after = clean_text(cell_text[end:end + 50])
                    results.append({
                        "id": blank_id,
                        "answer": answer,
                        "context_before": before,
                        "context_after": after,
                        "context_paragraph_idx": ti,
                    })
                    blank_id += 1

    return results


def extract_video_links(tables: list, paragraphs: list) -> list[dict]:
    """影片連結：先查表格，再 fallback 到段落純文字"""
    links = []

    # 表格查找
    for t in tables:
        if not t.rows:
            continue
        first_cell = clean_text(t.rows[0].cells[0].text) if t.rows[0].cells else ""
        if "影片連結" in first_cell or "影片" in first_cell:
            for row in t.rows:
                for ci, cell in enumerate(row.cells):
                    if ci == 0:
                        continue
                    text = clean_text(cell.text)
                    if not text:
                        continue
                    lines = [l.strip() for l in cell.text.split("\n") if l.strip()]
                    if not lines:
                        continue
                    title = clean_text(lines[0])
                    url = ""
                    for line in lines:
                        if "http" in line:
                            url = clean_text(line)
                            break
                    # Fallback: extract URL from hyperlinks in cell XML
                    if not url:
                        try:
                            xml_str = cell._tc.xml
                            url_match = re.search(r'(https?://[^\s"<>]+)', xml_str)
                            if url_match:
                                candidate = url_match.group(1)
                                # Skip XML namespace URIs
                                if ("schemas.openxmlformats.org" not in candidate
                                        and "schemas.microsoft.com" not in candidate):
                                    url = candidate
                        except Exception:
                            pass
                    if title:
                        links.append({"title": title, "url": url})
            if links:
                return links

    # Paragraph fallback: scan for URLs near "影片" keywords
    for i, p in enumerate(paragraphs):
        text = clean_text(p.text)
        if re.search(r"影片|video|youtu|youtube", text, re.IGNORECASE):
            url_match = re.search(r"(https?://[^\s<>\"]+)", text)
            if url_match:
                url = url_match.group(1)
                # Skip XML namespace URIs
                if "schemas.openxmlformats.org" not in url and "schemas.microsoft.com" not in url:
                    links.append({"title": text.split("http")[0].strip() or "影片", "url": url})
            # Also check next few paragraphs for URLs
            for j in range(i + 1, min(i + 4, len(paragraphs))):
                next_text = clean_text(paragraphs[j].text)
                url_match2 = re.search(r"(https?://[^\s<>\"]+)", next_text)
                if url_match2:
                    url2 = url_match2.group(1)
                    if "schemas.openxmlformats.org" not in url2 and "schemas.microsoft.com" not in url2:
                        links.append({"title": text[:50], "url": url2})
                        break

    return links


def extract_vocabulary(paragraphs: list, tables: list) -> list[dict]:
    """
    解析生字/詞語
    策略1: 樣式名含「填空題目」的段落 (G6 原有)
    策略2: 「(N) 詞語：定義」格式的任意段落
    策略3: 從「詞語題庫」表格
    """
    vocab = []
    seen_words = set()

    # 策略1: 教材：(1)＿：填空題目 樣式 (use flexible spacing around colon)
    pattern1 = re.compile(r"^\((\d+)\)\s+(.+?)\s*[:：]\s*(.+)$")
    for p in paragraphs:
        style_name = p.style.name if p.style else ""
        text = clean_text(p.text)
        if "填空題目" in style_name or "vocab" in style_name.lower():
            m = pattern1.match(text)
            if m:
                word = clean_text(m.group(2))
                definition = clean_text(m.group(3))
                if word and word not in seen_words:
                    vocab.append({"word": word, "definition": definition})
                    seen_words.add(word)

    # 策略2: (N) 詞語：定義（G4, G5 等）—  also covers 填空題目 style fallback
    if not vocab:
        for p in paragraphs:
            text = clean_text(p.text)
            m = pattern1.match(text)
            if m:
                word = clean_text(m.group(2))
                definition = clean_text(m.group(3))
                # 過濾非詞語的誤匹配（定義超短或詞太長）
                if word and len(word) <= 12 and len(definition) > 2 and word not in seen_words:
                    vocab.append({"word": word, "definition": definition})
                    seen_words.add(word)

    # 策略3: 從表格找「N. 詞語」格式（詞語題庫）
    if not vocab:
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    txt = clean_text(cell.text)
                    lines = [l.strip() for l in txt.split("\n") if l.strip()]
                    for line in lines:
                        m = re.match(r"^[A-J１-９1-9][.．]\s*(.{1,8})$", line)
                        if m:
                            word = clean_text(m.group(1))
                            if word and word not in seen_words:
                                vocab.append({"word": word, "definition": ""})
                                seen_words.add(word)

    return vocab


def extract_reading_strategy(paragraphs: list, tables: list, filename: str) -> tuple[str, str]:
    """
    從檔名和內容判斷 reading_strategy 和 reading_strategy_type
    Returns: (strategy_label, strategy_type)
    """
    # 從檔名括號中抽取策略名
    m = re.search(r"[（(]([^）)]+)[）)]", filename)
    strategy_label = clean_text(m.group(1)) if m else ""

    # 判斷 type
    strategy_type = "general"
    label_lower = strategy_label.lower()

    if "圖文" in strategy_label:
        strategy_type = "graphic_text_integration"
    elif "摘要" in strategy_label and "問題" in strategy_label and "解決" in strategy_label:
        strategy_type = "summary_psr"
    elif "摘要" in strategy_label:
        strategy_type = "summary"
    elif "推論" in strategy_label:
        strategy_type = "inference"
    elif "解決問題" in strategy_label or "科學探究" in strategy_label:
        strategy_type = "problem_solving"
    elif "比較" in strategy_label:
        strategy_type = "compare_contrast"
    elif "自我提問" in strategy_label:
        strategy_type = "self_questioning"
    elif "寫作手法" in strategy_label:
        strategy_type = "writing_technique"
    elif "文言文" in strategy_label or filename.startswith("文-"):
        strategy_type = "classical_chinese"

    return strategy_label, strategy_type


def extract_self_check_items(paragraphs: list) -> list[str]:
    """自我檢核清單"""
    items = []
    in_check = False
    for p in paragraphs:
        text = clean_text(p.text)
        if "自我檢核" in text:
            in_check = True
            continue
        if in_check:
            if re.match(r"^[□☐✓✗]?\s*\d+[.．]\s*.+", text):
                item = re.sub(r"^[□☐✓✗]?\s*\d+[.．]\s*", "", text).strip()
                if item:
                    items.append(item)
            elif text.startswith("◎") or text.startswith("※"):
                in_check = False
    return items


def extract_discussion_questions(paragraphs: list) -> list[str]:
    """小試身手 / 進階挑戰"""
    items = []
    in_section = False
    for p in paragraphs:
        text = clean_text(p.text)
        if "小試身手" in text or "進階挑戰" in text or "思考題" in text:
            in_section = True
            continue
        if in_section:
            if text.startswith("◎") or text.startswith("※") or "自我檢核" in text:
                in_section = False
            elif text:
                items.append(text)
    return items


def extract_multiple_choice(paragraphs: list) -> list[dict]:
    """選擇題"""
    def extract_answer(txt):
        m = re.search(r"[（(]\s*([A-Da-dＡ-Ｄ])\s*[）)]", txt)
        if m:
            return m.group(1).upper().translate(str.maketrans("ＡＢＣＤ", "ABCD"))
        return ""

    questions = []
    i = 0
    while i < len(paragraphs):
        text = clean_text(paragraphs[i].text)
        ans = extract_answer(text)
        q_match = re.match(r"^[（(]\s*[A-Da-dＡ-Ｄ]\s*[）)]\s*\d+[.．]", text)
        if q_match and ans:
            question_text = re.sub(r"^[（(]\s*[A-Za-z]\s*[）)]\s*", "", text).strip()
            options = []
            i += 1
            while i < len(paragraphs):
                opt = clean_text(paragraphs[i].text)
                if re.match(r"^[A-Da-d][.．]", opt):
                    parts = re.split(r"\s{2,}(?=[A-Da-d][.．])", opt)
                    for part in parts:
                        if re.match(r"^[A-Da-d][.．]", part.strip()):
                            options.append(re.sub(r"^[A-Da-d][.．]\s*", "", part.strip()))
                    i += 1
                elif re.match(r"^[（(]\s*[A-Da-d]\s*[）)]\s*\d+", opt):
                    break
                else:
                    i += 1
                    break
            questions.append({"question": question_text, "options": options, "answer": ans})
            continue
        i += 1
    return questions


def extract_story_structure_table(tables: list) -> Optional[list]:
    """問題-解決-結果表格 (支援多種格式)

    格式一 (G6-L22, PSR 標題行):
        Row 0: ['課文標題', '課文標題', ...]
        Row 1: ['問題', '問題', '內容']
        Row 2: ['解決', '解決', '內容']
        → Row 0/1 各有 ≥2 個 PSR keyword

    格式二 (G6-L24/25, 三欄 header):
        Row 0: ['元素/段落', '提示', '重點']
        Row 1: ['問題/2', '描述提示', '重點內容']
        Row 2: ['解決/3~6', '描述提示', '重點內容']
        → Row 0 有 '元素' 或 '段落'，Row 1+ 的 col0 包含 PSR keyword

    格式三 (G6-L23, 課文標題 + PSR 行):
        Row 0: ['課文標題', '課文標題']
        Row 1: ['問題', '...']
        Row 2: ['解決 ...', '...']
        → col0 中有 ≥2 行含 PSR keyword
    """
    # PSR keywords that appear in cell0 of content rows
    psr_keywords = ["問題", "解決", "結果", "迴響", "研究問題", "假說"]
    # Header keywords that indicate a 3-column structured table (format 2)
    header_keywords = ["元素", "段落"]

    def _extract_rows(t) -> list:
        rows_data = []
        for row in t.rows:
            cells_clean = []
            for c in row.cells:
                txt = clean_text(c.text)
                if not cells_clean or txt != cells_clean[-1]:
                    cells_clean.append(txt)
            if any(cells_clean):
                rows_data.append(cells_clean)
        return rows_data

    for t in tables:
        if not t.rows:
            continue

        # --- Format 1: classic PSR (≥2 PSR keywords in row 0 or row 1) ---
        for ri in range(min(2, len(t.rows))):
            cells = [clean_text(c.text) for c in t.rows[ri].cells]
            if sum(1 for c in cells if any(kw in c for kw in psr_keywords)) >= 2:
                return _extract_rows(t)

        # --- Format 2: 3-col header table (元素/段落 in row 0) ---
        if len(t.rows) >= 2:
            row0_cells = [clean_text(c.text) for c in t.rows[0].cells]
            row0_flat = " ".join(row0_cells)
            if any(hkw in row0_flat for hkw in header_keywords):
                # Verify row 1+ col0 has PSR content
                psr_count = 0
                for ri in range(1, len(t.rows)):
                    cell0 = clean_text(t.rows[ri].cells[0].text) if t.rows[ri].cells else ""
                    if any(kw in cell0 for kw in psr_keywords):
                        psr_count += 1
                if psr_count >= 1:
                    return _extract_rows(t)

        # --- Format 3: 2-col table with title row + PSR content rows ---
        # e.g. G6-L23: Row0=['課文標題'], Row1=['問題','...'], Row2=['解決 ...','...']
        # Only ≥1 PSR keyword in row0/row1 but col0 of multiple rows has PSR keywords
        if len(t.rows) >= 3:
            psr_col0_count = 0
            for ri in range(len(t.rows)):
                cell0 = clean_text(t.rows[ri].cells[0].text) if t.rows[ri].cells else ""
                if any(kw in cell0 for kw in psr_keywords):
                    psr_col0_count += 1
            if psr_col0_count >= 2:
                return _extract_rows(t)

    return None


def extract_graphic_text_structure(paras: list, strategy_label: str) -> Optional[list]:
    """圖文整合策略課文的結構表 — 從段落步驟文字合成 story_structure_table.

    適用於 G7-L29/30 圖文整合閱讀策略（docx 內沒有 PSR 表格），
    提取 ❶❷❸❹ 步驟說明合成結構資料，讓 fast path 可用。
    """
    if "圖文" not in strategy_label:
        return None

    step_pattern = re.compile(r"^[❶❷❸❹①②③④]\s*(.+)$")
    steps = []
    for p in paras:
        text = clean_text(p.text)
        m = step_pattern.match(text)
        if m:
            steps.append(text)

    if len(steps) >= 3:
        return [["圖文整合閱讀步驟"]] + [[s] for s in steps]
    return None


def extract_reading_benchmark(tables: list) -> dict:
    """朗讀速率標準"""
    for t in tables:
        cells_flat = [clean_text(c.text) for row in t.rows for c in row.cells]
        if any("還要多加練習" in c for c in cells_flat) or any("朗讀速率" in c for c in cells_flat):
            if len(t.rows) >= 2:
                thresholds = [clean_text(c.text) for c in t.rows[0].cells]
                feedbacks = [clean_text(c.text) for c in t.rows[1].cells]
                levels = []
                for t_val, f_val in zip(thresholds, feedbacks):
                    if t_val and f_val:
                        levels.append({"threshold": t_val, "feedback": f_val})
                if levels:
                    return {"levels": levels}
    return {}


# ---------- image extraction (improved) ----------

def extract_images_with_metadata(
    doc: docx.Document, output_dir: Path, lesson_code: str
) -> list[dict]:
    """
    Extract embedded images with metadata:
      - filename
      - size_bytes
      - image_hash (sha256[:12])
      - content_type
    Note: paragraph_index not available via relationships API;
    we use relationship order as proxy.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    image_list = []
    seen_hashes = set()

    for rel in doc.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            img_data = rel.target_part.blob
            img_hash = hashlib.sha256(img_data).hexdigest()[:12]

            # Skip exact duplicates
            if img_hash in seen_hashes:
                continue
            seen_hashes.add(img_hash)

            content_type = rel.target_part.content_type
            ext = content_type.split("/")[-1]
            if ext == "jpeg":
                ext = "jpg"
            elif ext == "vnd.openxmlformats-officedocument.drawing":
                ext = "png"

            idx = len(image_list) + 1
            filename = f"{lesson_code}-{idx:02d}.{ext}"
            out_path = output_dir / filename

            with open(out_path, "wb") as f:
                f.write(img_data)

            image_list.append({
                "filename": f"images/{lesson_code}/{filename}",
                "size_bytes": len(img_data),
                "image_hash": img_hash,
                "content_type": content_type,
                "caption": "",
            })
        except Exception as e:
            # Log but continue
            pass

    return image_list


# ---------- universal parser ----------

def parse_lesson(doc_path: Path, lesson_code: str, images_base: Path) -> tuple[dict, list[str]]:
    """
    Universal lesson parser — works for G4~G9 and 文言文.
    Returns (data_dict, warnings_list)
    """
    warnings = []
    doc = docx.Document(str(doc_path))
    paras = doc.paragraphs
    tables = doc.tables

    # Determine grade number
    grade_str = lesson_code.split("-")[0]  # "G4", "G5", ... "文"
    if grade_str.startswith("G") and grade_str[1:].isdigit():
        grade_num = int(grade_str[1:])
    else:
        grade_num = 0  # 文言文

    # Reading strategy from filename
    strategy_label, strategy_type = extract_reading_strategy(paras, tables, doc_path.name)

    result = {
        "lesson_code": lesson_code,
        "grade": grade_num,
        "grade_code": lesson_code,
        "reading_strategy": strategy_label,
        "reading_strategy_type": strategy_type,
        "source_file": doc_path.name,
        "parsed_at": datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
        "flags": [],
    }

    # Genre heuristic
    if "文言文" in doc_path.name or lesson_code.startswith("文-"):
        result["genre"] = "文言文"
    elif any(kw in doc_path.name for kw in ["說明", "科學", "報導", "新聞", "議論"]):
        result["genre"] = "說明文"
    elif any(kw in doc_path.name for kw in ["感情", "日記", "情緒", "小說"]):
        result["genre"] = "記敘文"
    else:
        result["genre"] = "說明文"  # default

    # --- Table 0: 課文主表 ---
    # G7 特殊：課文在 tables[1]（tables[0] 是說明框）
    table_idx = 0
    if len(tables) > 1 and grade_num == 7:
        # Check if table[0] has lesson title
        row0_cells = [clean_text(c.text) for c in tables[0].rows[0].cells] if tables[0].rows else []
        if not any(re.match(r"第\d+課", c) for c in row0_cells):
            table_idx = 1

    if tables and table_idx < len(tables):
        title, authors, story_text = get_title_from_table0(tables[table_idx])
        result["title"] = title
        result["authors"] = authors
        result["story_text"] = story_text

        if not title:
            warnings.append(f"title_empty: Table {table_idx} Row 0 no '第N課' pattern")
            result["flags"].append("title_empty")

        paras_list = split_story_paragraphs(story_text)
        # #2218: 剝離「圖N …／表N …」圖說行，避免被當成課文段落 render。
        # caption_rows 稍後 (--- 圖片 extract ---) 寫進 images[].caption。
        paras_list, caption_rows = split_caption_rows(paras_list)
        result["_caption_rows"] = caption_rows  # internal; popped before save
        result["paragraphs"] = paras_list
        result["paragraph_count"] = len(paras_list)
        if caption_rows:
            # 有圖說行被剝離時，story_text / char_count 須與 paragraphs 同步，
            # 否則 story_text 仍含被移除的圖說。無圖說行時維持原行為（不動）。
            result["story_text"] = "\n".join(paras_list)
            result["char_count"] = sum(len(p) for p in paras_list)
        else:
            result["char_count"] = len(story_text.replace(" ", ""))
    else:
        result["title"] = ""
        result["authors"] = ""
        result["story_text"] = ""
        result["paragraphs"] = []
        result["paragraph_count"] = 0
        result["char_count"] = 0
        warnings.append("no_tables: Cannot extract story text")
        result["flags"].append("no_tables")

    # --- 填空題 (fill_in_blank) ---
    fill_in_blank = extract_fill_in_blank(tables, story_table_idx=table_idx)
    result["fill_in_blank"] = fill_in_blank
    result["fill_in_blank_count"] = len(fill_in_blank)

    # --- 生字 ---
    vocab = extract_vocabulary(paras, tables)
    result["vocabulary"] = vocab
    result["vocabulary_count"] = len(vocab)
    if not vocab:
        warnings.append("vocab_empty: No vocabulary extracted")
        result["flags"].append("vocab_empty")

    # --- 朗讀速率表 ---
    reading_benchmark = extract_reading_benchmark(tables)
    if reading_benchmark:
        result["reading_benchmark"] = reading_benchmark

    # --- 影片連結 ---
    video_links = extract_video_links(tables, paras)
    result["video_links"] = video_links

    # --- 選擇題 ---
    mc = extract_multiple_choice(paras)
    result["multiple_choice"] = mc
    result["multiple_choice_count"] = len(mc)

    # --- 問題-解決-結果表格 (含圖文整合 fallback) ---
    story_struct = extract_story_structure_table(tables)
    if not story_struct:
        story_struct = extract_graphic_text_structure(paras, strategy_label)
    if story_struct:
        result["story_structure_table"] = story_struct

    # --- 自我檢核 ---
    self_check = extract_self_check_items(paras)
    result["self_check_items"] = self_check

    # --- 小試身手 ---
    discussion = extract_discussion_questions(paras)
    result["discussion_content"] = discussion

    # --- G7 策略練習 ---
    if grade_num == 7 or "圖文" in strategy_label:
        strategy_exercises = []
        current_exercise = None
        for p in paras:
            text = clean_text(p.text)
            m_ex = re.match(r"^(練習[一二三四五六])[:：]?\s*(.*)$", text)
            if m_ex:
                if current_exercise:
                    strategy_exercises.append(current_exercise)
                current_exercise = {
                    "exercise": m_ex.group(1),
                    "description": clean_text(m_ex.group(2)),
                    "steps": [],
                }
                continue
            m_step = re.match(r"^(步驟[❶❷❸❹①②③④])[:：]?\s*(.+)$", text)
            if m_step and current_exercise is not None:
                current_exercise["steps"].append({
                    "step": m_step.group(1),
                    "description": m_step.group(2),
                })
        if current_exercise:
            strategy_exercises.append(current_exercise)
        result["strategy_exercises"] = strategy_exercises

    # --- 圖片 extract (with metadata) ---
    img_output_dir = images_base / lesson_code
    images = extract_images_with_metadata(doc, img_output_dir, lesson_code)
    # #2218: 把先前剝離的「圖N」圖說寫進對應 image 的 caption。
    caption_rows = result.pop("_caption_rows", [])
    assign_captions_to_images(images, caption_rows)
    result["images"] = images
    result["image_count"] = len(images)

    return result, warnings


# ---------- YAML helpers ----------

def represent_str(dumper, data):
    if "\n" in data or len(data) > 80:
        return dumper.represent_scalar("tag:yaml.org,2002:str", data, style="|")
    return dumper.represent_scalar("tag:yaml.org,2002:str", data)


yaml.add_representer(str, represent_str)


def save_yaml(data: dict, path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        yaml.dump(data, f, allow_unicode=True, default_flow_style=False, sort_keys=False, width=120)


# ---------- file discovery ----------

def discover_docx_files(source_base: Path) -> list[tuple[Path, str]]:
    """
    Discover all teacher-version docx files.
    Returns list of (docx_path, lesson_code)
    """
    results = []
    seen_codes = {}

    def scan_dir(d: Path):
        for item in sorted(d.iterdir()):
            if item.is_dir():
                scan_dir(item)
            elif item.suffix.lower() == ".docx" and not item.name.startswith("~$"):
                code = extract_lesson_code_from_filename(item.name)
                if code:
                    if code in seen_codes:
                        # Duplicate: prefer the one in 教師版 dir
                        prev_path = seen_codes[code][0]
                        if TEACHER_DIR_NAME in str(item) and TEACHER_DIR_NAME not in str(prev_path):
                            seen_codes[code] = (item, code)
                        # else keep existing
                    else:
                        seen_codes[code] = (item, code)
                else:
                    print(f"WARN: cannot extract lesson code from: {item.name}", file=sys.stderr)

    # Scan teacher version (L1-122)
    teacher_dir = source_base / TEACHER_DIR_NAME
    if teacher_dir.exists():
        scan_dir(teacher_dir)
    else:
        print(f"WARN: Teacher dir not found: {teacher_dir}", file=sys.stderr)

    # Scan Stage 3 (L123-L157)
    stage3_dir = source_base / STAGE3_DIR_NAME
    if stage3_dir.exists():
        scan_dir(stage3_dir)
    else:
        print(f"WARN: Stage3 dir not found: {stage3_dir}", file=sys.stderr)

    results = sorted(seen_codes.values(), key=lambda x: x[1])
    return results


# ---------- main ----------

def main():
    parser = argparse.ArgumentParser(description="批量解析 158 課教師版 docx → YAML")
    parser.add_argument(
        "--source",
        default="private/curriculum-source/2026-05-01/1.L1-158新版完成學習單1150415",
        help="Source directory containing 1-1教師版 and 第三階段 subdirectories",
    )
    parser.add_argument(
        "--output",
        default="backend/data/lessons/_parsed_2026-05-01",
        help="Output directory for YAML files",
    )
    parser.add_argument(
        "--images",
        default="backend/data/lessons/_parsed_2026-05-01/images",
        help="Output directory for extracted images",
    )
    parser.add_argument(
        "--force", action="store_true", help="Re-parse even if output YAML already exists"
    )
    parser.add_argument(
        "--lesson", default=None, help="Only parse specific lesson code (e.g. G4-L1)"
    )
    parser.add_argument(
        "--filter", default=None, dest="filter_re",
        help="Only parse lesson codes matching this regex (e.g. 'G6-L2[2-5]|G7-L(28|29|30)')"
    )
    parser.add_argument(
        "--confirm", action="store_true",
        help="Required when using --filter to confirm selective re-parse"
    )
    args = parser.parse_args()

    source_dir = Path(args.source)
    output_dir = Path(args.output)
    images_dir = Path(args.images)

    if not source_dir.exists():
        print(f"ERROR: source dir not found: {source_dir}", file=sys.stderr)
        sys.exit(1)

    output_dir.mkdir(parents=True, exist_ok=True)
    images_dir.mkdir(parents=True, exist_ok=True)

    print(f"Scanning: {source_dir}")
    all_files = discover_docx_files(source_dir)
    print(f"Found {len(all_files)} docx files")

    if args.lesson:
        all_files = [(p, c) for p, c in all_files if c == args.lesson]
        if not all_files:
            print(f"ERROR: lesson {args.lesson} not found", file=sys.stderr)
            sys.exit(1)

    if args.filter_re:
        if not args.confirm:
            print(f"ERROR: --filter requires --confirm to prevent accidental bulk re-parse", file=sys.stderr)
            sys.exit(1)
        try:
            filter_pattern = re.compile(args.filter_re)
        except re.error as e:
            print(f"ERROR: invalid --filter regex: {e}", file=sys.stderr)
            sys.exit(1)
        all_files = [(p, c) for p, c in all_files if filter_pattern.search(c)]
        if not all_files:
            print(f"ERROR: no lessons match filter '{args.filter_re}'", file=sys.stderr)
            sys.exit(1)
        print(f"Filter '{args.filter_re}' matched {len(all_files)} lessons: {[c for _, c in all_files]}")

    # Tracking
    success = []
    partial = []
    failed = []
    skipped = []

    start_time = datetime.now()

    for i, (doc_path, lesson_code) in enumerate(all_files):
        out_yaml = output_dir / f"{lesson_code}.yml"

        # Idempotent: skip if exists (unless --force)
        if out_yaml.exists() and not args.force:
            skipped.append(lesson_code)
            continue

        elapsed = (datetime.now() - start_time).total_seconds()
        if elapsed > 25 * 60:  # 25 min soft limit
            print(f"\nWARN: 25 min limit reached at lesson {i+1}/{len(all_files)}. Stopping.")
            break

        print(f"[{i+1}/{len(all_files)}] {lesson_code} ... ", end="", flush=True)

        try:
            data, warnings = parse_lesson(doc_path, lesson_code, images_dir)
            save_yaml(data, out_yaml)

            img_count = data.get("image_count", 0)
            warn_str = f" [{', '.join(warnings[:2])}]" if warnings else ""
            print(f"OK (title={data.get('title', '')[:20]!r}, imgs={img_count}){warn_str}")

            if warnings:
                partial.append((lesson_code, warnings))
            else:
                success.append(lesson_code)

        except Exception as e:
            tb = traceback.format_exc()
            print(f"FAIL: {e}")
            failed.append((lesson_code, str(e), tb))

    # --- Write parse-report.md ---
    all_processed_codes = success + [c for c, _ in partial] + skipped
    total_images = sum(
        len(list((images_dir / code).glob("*")))
        for code in all_processed_codes
        if (images_dir / code).exists()
    )

    report_lines = [
        "# Parse Report — 158 課批量解析",
        f"\n生成時間: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## Summary",
        "",
        f"| 狀態 | 數量 |",
        f"|------|------|",
        f"| Success (clean) | {len(success)} |",
        f"| Partial (warnings) | {len(partial)} |",
        f"| Failed | {len(failed)} |",
        f"| Skipped (already exist) | {len(skipped)} |",
        f"| Total processed | {len(success) + len(partial) + len(failed)} |",
        f"| Total images extracted | {total_images} |",
        "",
        "## Partial Parses (有 warnings 但成功輸出)",
        "",
    ]

    for code, warns in sorted(partial):
        report_lines.append(f"### {code}")
        for w in warns:
            report_lines.append(f"- {w}")
        report_lines.append("")

    report_lines += [
        "## Failed Parses",
        "",
    ]

    for code, err, tb in sorted(failed):
        report_lines.append(f"### {code}")
        report_lines.append(f"```")
        report_lines.append(f"Error: {err}")
        # Show first 5 lines of traceback
        tb_lines = tb.strip().split("\n")
        for line in tb_lines[-5:]:
            report_lines.append(line)
        report_lines.append(f"```")
        report_lines.append("")

    report_lines += [
        "## Known Remaining Gaps",
        "",
        "- Multi-lesson files (e.g. G4-L20-22, G9-L15~16) only parse first lesson metadata",
        "  → Follow-up: detect multi-lesson files and emit N YAMLs per file",
        "- `paragraph_index` for images not available via rels API (would need XML walk)",
        "  → Current metadata: filename, size_bytes, image_hash, content_type",
        "- Some vocab fields may be empty for lessons with non-standard styles",
        "  → Check partial parse list above for vocab_empty flags",
        "- Video URLs blank where embedded as Word hyperlinks (not plain text)",
        "  → Fallback extracts from paragraph text; hyperlink XML fallback attempted",
        "",
        "## Follow-up PRs (do not include in this PR)",
        "",
        "- Schema-match against L01.yml format (curriculum index integration)",
        "- Filter decorative-chrome images (#1341) using image_hash dedup",
        "- Wire 158 lessons into platform routing (#1344 part 2)",
        "- Multi-lesson file splitting (emit N YAMLs for L15~16 etc.)",
    ]

    report_path = output_dir / "parse-report.md"
    report_path.write_text("\n".join(report_lines), encoding="utf-8")

    # --- Final summary ---
    print("\n" + "=" * 70)
    print("PARSE COMPLETE")
    print("=" * 70)
    print(f"Success (clean):   {len(success)}")
    print(f"Partial (warnings): {len(partial)}")
    print(f"Failed:            {len(failed)}")
    print(f"Skipped (cached):  {len(skipped)}")
    print(f"Total images:      {total_images}")
    print(f"Report:            {report_path}")
    print("=" * 70)

    if failed:
        print("\nFailed lessons:")
        for code, err, _ in failed:
            print(f"  {code}: {err}")

    return 0 if not failed else 1


if __name__ == "__main__":
    sys.exit(main())
