#!/usr/bin/env python3
"""重造測試用的合成學習單（#2865）。

⛔ 內容全是自己編的假課文 —— 真原稿在 gitignore 的 private/，不能進 repo。
但**版面結構跟真的一樣**，那才是見證對帳要測的東西。

    python3 backend/tests/fixtures/build_synthetic_worksheet.py
"""
import pathlib
import subprocess

from docx import Document
from docx.shared import Pt

HERE = pathlib.Path(__file__).resolve().parent
REPO = HERE.parents[2]


def main() -> None:
    d = Document()

    def para(t, size=11, bold=False):
        p = d.add_paragraph()
        r = p.add_run(t)
        r.font.size = Pt(size)
        r.bold = bold

    para("一        讀全文-做記號", 14, True)
    para("這是第一大題的說明文字，請讀完全文並做記號。")
    for i in range(1, 4):
        para(f"({i}) 這是第一大題的第 {i} 小題。")
    d.add_page_break()
    # 續頁刻意不再印標題 —— 真學習單就是這樣，而第一版的裁判因此整頁跳過
    para("（第一大題續，這一頁沒有再印一次標題）")
    para("(4) 跨頁之後的第 4 小題。")
    para("二        念順順", 14, True)
    para("請把下面這段念順。")
    d.add_page_break()
    # 標題後面還有字 —— 真學習單也這樣，而第一版的 `\s*$` 對不上
    para("三        語詞我最棒        在空格內填入語詞", 14, True)
    para("本課語詞：甲乙、丙丁、戊己")
    for i, w in enumerate(["甲乙", "丙丁", "戊己"], 1):
        para(f"({i})  {w}  ：這是第 {i} 個語詞的解釋。")
    # 同一頁上的隔壁節 —— 不分節就會被算進來
    para("四        語詞應用", 14, True)
    for i in range(1, 3):
        para(f"({i}) 這是語詞應用的第 {i} 小題，不該被算進語詞我最棒。")

    docx = HERE / "synthetic_worksheet.docx"
    d.save(docx)
    subprocess.run(
        ["bash", str(REPO / "scripts" / "docx_to_pdf.sh"), str(docx), str(HERE), "synthetic"],
        check=True,
    )
    (HERE / "synthetic.pdf").rename(HERE / "synthetic_worksheet.pdf")
    print(f"✅ {docx.name} + synthetic_worksheet.pdf")


if __name__ == "__main__":
    main()
