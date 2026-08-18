#!/usr/bin/env python3
"""extract_key_reading.py — 重點朗讀（念順順）指定段落, from section 二 (#2683, #2720).

WHY THIS EXISTS
---------------
Every one of the 175 lessons used to fall back to reading the whole text aloud, which
is the exact thing the 2026-07-20 expert review ruled against: 朗讀只練老師指定的重點段,
不練全文.

The first edition had this data. It is not reusable as an answer key — it is keyed by
lesson code, the second edition renumbered every lesson, and 23 of the 53 comparable
lessons had their text rewritten. It is a regression golden set, not a source.

WHERE THE PASSAGE ACTUALLY LIVES
--------------------------------
Section 二 念順順 does not contain the passage. It names it:

    請用計時器，從指定段落（三）開始朗讀，計時 1 分鐘讀的字數…

So the anchor is a paragraph NUMBER. The question this file has to answer correctly is
what that number counts.

THE BUG THIS REPLACES, AND WHY IT WAS INVISIBLE (#2720)
-------------------------------------------------------
It used to answer with `body[anchor - 1]` against `body.yml` — the paragraph list that
`extract_lesson_body` DERIVES with a chain of heuristics (a 46-character floor, a
short-paragraph recovery capped at three, chrome prefixes, instruction markers,
exercise marks, dedup, five tiers of boundary heading).

「第三段」 does not mean "the third surviving element of that list". It means the
paragraph the worksheet PRINTS as 三. Those two coincide only when no heuristic
happened to drop or keep a paragraph ahead of the anchor, and 靖杭's scan of the
first edition's 148 hand-checked passages measured how often that holds:

    課文沒改、段界一致的 28 課 → 錨點只有 17 課正確，錯的 11 課有 10 課差一整段

Worse, the check that was supposed to catch this reported `confirmed` on 8 of those 11.
It took the first edition's passage TEXT, used it only to look up a paragraph NUMBER,
and then compared two numbers from two different coordinate systems. Comparing the text
it already held would have caught every one.

WHERE THE NUMBERING ACTUALLY IS
-------------------------------
The worksheet prints its own paragraph numbers, in a column of the body table:

    row: [ 段號欄 ] [ 課文欄 ] [ 字數欄 ]
          ㄧ⏎二⏎⏎⏎三⏎⏎⏎四…   最近心情起起伏伏…⏎上課時…⏎然而最近…   27⏎57⏎85⏎…

So the answer is read, not re-derived: the anchor indexes the 課文 CELL's own
paragraphs. On 靖杭's golden set that moves the hit rate from 20/34 to 30/34, and the
four remaining misses are all lessons where the printed numbering and the cell's
paragraph count disagree — which is a condition this file can test.

Two traps in that column, both real:

  · the first mark is often 「ㄧ」 = U+3127 BOPOMOFO LETTER I, not 「一」 = U+4E00. A
    CJK-only numeral class does not match it, and the lesson then looks unnumbered.
  · python-docx reports a horizontally merged cell once per grid column, so the 課文
    text appears two or three times in one row. Dedup before picking the longest.

THE CHECKS, AND WHAT THEY COST
------------------------------
Nothing is written unless the numbering can be trusted, because a student reading the
wrong paragraph aloud is worse than a student reading the whole text — the whole-text
fallback is degraded, a confidently wrong passage is a lie.

  1. The 課文 cell must hold as many paragraphs as the author numbered, give or take a
     couple of unnumbered tail lines. Measured over all 175 second-edition worksheets,
     against the 33 lessons 靖杭's golden set can judge:

         paras - marks     correct   wrong
              0               22        2      L0030, L0072
              1                5        1      L0110
              2                2        0
              3, 5, 6, 10      0        4      L0122, L0046, L0071, L0060

     One or two unnumbered paragraphs are closing lines the author did not number, and
     they sit after the anchor, so indexing still lands — 2 recovers seven lessons over
     strict equality, both judgeable ones correct. Three or more means the cell and the
     numbering disagree about where paragraphs BEGIN, and every such lesson was wrong.
     `paras < marks` fails too: the cell lost a numbered paragraph.

     This gate alone runs ~90% (29/32). It is not the last line of defence — the three
     that slip past it are exactly the three check 2 catches, so nothing wrong reaches a
     lesson that has a first-edition counterpart. Lessons without one rest on this gate,
     which is why the CI golden test exists and why the number is stated rather than
     rounded up.
  2. The first edition's PASSAGE TEXT, when this lesson has one. Compared as text
     against the passage about to be written, not as a paragraph number.
  3. The passage must be one of `body.yml`'s paragraphs. It is what the student reads,
     what the TTS sentence table is built from, and what #2718's test asserts; a
     passage outside that list means the body extractor and the numbering disagree
     about paragraph boundaries (Word splits a paragraph at a manual line break in a
     handful of lessons), and the tail would be served missing.

Two checks that were tried and do NOT work, recorded so they are not tried again:

  · 字數欄 as a completeness check. Its max runs 280–520 while the body runs 535–1670
    characters, because the column counts the printed one-minute excerpt, not the text.
    (This also retires the first edition's belief that max 累計字數 = the passage
    length — the rule that produced the 370-character passages of #2712.)
  · re-deriving the numbering by counting prose paragraphs across the whole document.
    Tried; 22/30 against the golden set, because other table columns, teacher
    annotations and Word's duplicated runs all land in the same paragraph stream.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import unicodedata
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "scripts"))

from extract_lesson_body import _paragraphs, extract as extract_body  # noqa: E402

LEGACY_TABLE = ROOT / "backend" / "data" / "key_reading_passages.yml"

#: 「ㄧ」 U+3127 is BOPOMOFO LETTER I, not the CJK 一. It is what the worksheet actually
#: prints for the first paragraph in a large share of lessons, and leaving it out of the
#: numeral class made those lessons read as having no numbering at all.
_BOPOMOFO_ONE = "ㄧ"
_CN_UNITS = {c: i for i, c in enumerate("一二三四五六七八九十", start=1)}
_NUMERAL_CHARS = "一二三四五六七八九十" + _BOPOMOFO_ONE

#: 「從指定段落（四）開始朗讀」 — the parenthesis is sometimes half-width, sometimes
#: unclosed, and sometimes padded with the tab that follows it in the form.
_ANCHOR = re.compile(
    r"從指定段落\s*[（(]?\s*([" + _NUMERAL_CHARS + r"]+|\d+)\s*[)）]?"
)
_ANCHOR_LOOSE = re.compile(
    r"指定段落[^0-9" + _NUMERAL_CHARS + r"]{0,6}([" + _NUMERAL_CHARS + r"]+|\d+)"
)

#: Bounds taken from what the professor actually marked, not from what seems reasonable.
#: The first edition's 134 marked ranges run 19 to 412 characters, median 147; six are
#: under 40. These exist only to catch a paragraph that is obviously not prose (a stray
#: caption, a whole page swallowed), so they sit outside the observed range.
MIN_CHARS, MAX_CHARS = 12, 900

#: How many paragraphs the 課文 cell may hold beyond the numbered ones. One or two are
#: closing lines the author left unnumbered, and they sit after the anchor. Three is
#: where the two disagree about paragraph STARTS. See the module docstring for the
#: measurement behind this bound.
MAX_UNNUMBERED_TAIL = 2

#: 說明文的小標題不算一段。The author numbers the prose paragraphs and not the
#: sub-headings above them — 《秋行軍蟲》 numbers six paragraphs while its cell holds
#: twelve, and the extras are 「人工噴農藥」「無人機噴農藥」「寄生蜂片」「寄生蜂扭蛋」.
#: A heading is short and does not end a sentence.
_HEADING_MAX = 14
#: 「……」 ends a paragraph. Leaving it out merged 《感情小日記2》's opening line into the
#: next paragraph and moved every anchor after it — the exact misfire #2726 warned this
#: signal has, caught on the lesson this whole mechanism was worked out on.
_SENTENCE_END = "。！？」』…⋯"

#: A 段號 cell holds only numerals, one per paragraph. Three is the smallest run that
#: cannot be a stray character — two would match a 「一/二」 heading pair.
_MIN_MARKS = 3
#: The 課文 cell has to actually hold the text. Below this it is a caption column.
_MIN_CELL_CHARS = 100


def cn_to_int(s: str) -> int | None:
    """Chinese or Arabic numeral → int.

    NFKC first: the worksheets mix full-width digits into the same field. 二十 and
    二十一 appear in the longer lessons and used to return None, which silently dropped
    those lessons rather than reporting anything.
    """
    s = unicodedata.normalize("NFKC", s.strip()).replace(_BOPOMOFO_ONE, "一")
    if s.isdigit():
        return int(s)
    if s in _CN_UNITS:
        return _CN_UNITS[s]
    if len(s) == 2 and s[0] == "十" and s[1] in _CN_UNITS:      # 十一 – 十九
        return 10 + _CN_UNITS[s[1]]
    if len(s) == 2 and s[1] == "十" and s[0] in _CN_UNITS:      # 二十, 三十 …
        return _CN_UNITS[s[0]] * 10
    if len(s) == 3 and s[1] == "十" and s[0] in _CN_UNITS and s[2] in _CN_UNITS:
        return _CN_UNITS[s[0]] * 10 + _CN_UNITS[s[2]]           # 二十一 …
    return None


#: 念順順 sets a fluency target, calibrated per lesson, printed as a run of checkbox
#: lines followed by the marker's message for each tier:
#:
#:      □ ＜190字        □ 191~220字      □ ＞221字          (146 lessons, characters/minute)
#:      還要多加練習。    嗯，還不錯喲！    哇嗚，超級厲害！
#:
#:      □42秒以下        □42~55秒         □55秒以上          (12 lessons, 文言文, seconds)
#:      速度有點快…      強者就是你啦！    速度再快一點！
#:
#: Stored as the raw worksheet strings under the `reading_benchmark` contract that
#: `frontend/src/utils/fluencyAnalyzer.ts` already parses — it handles both units and
#: knows that a seconds benchmark has no characters-per-minute pass mark. Emitting a
#: parsed {min,max} shape instead would have been a second contract for the same data,
#: and the frontend was already reading nothing: `lesson_indexes` hard-coded
#: `reading_benchmark: None`, so every lesson fell back to a grade-wide default while
#: its own thresholds sat in the DOCX.
_BOX = re.compile(r"^[□☐■◻▢]\s*")
_TIER = re.compile(r"\d{2,4}\s*[字秒]")


def find_reading_benchmark(paras: list[str]) -> dict | None:
    """The tiers and their messages, or None where the worksheet sets no target.

    A tier line is a checkbox carrying a number and a unit. They run consecutively; the
    same number of non-empty lines after them are the messages, paired by position —
    which is what the seconds form needs, since there the FIRST tier is the fast one and
    its message is a caution rather than praise.
    """
    text = [p.strip() for p in paras]
    i = 0
    while i < len(text):
        if not (_BOX.match(text[i]) and _TIER.search(text[i])):
            i += 1
            continue
        run = []
        while i < len(text) and _BOX.match(text[i]) and _TIER.search(text[i]):
            run.append(_BOX.sub("", text[i]).strip())
            i += 1
        if not 2 <= len(run) <= 4:
            continue
        msgs = [t for t in text[i:i + len(run) + 4] if t][:len(run)]
        if len(msgs) < len(run):
            msgs = [""] * len(run)
        return {"levels": [{"threshold": t, "feedback": m} for t, m in zip(run, msgs)]}
    return None


def find_anchor(paras: list[str]) -> int | None:
    text = "\n".join(paras)
    m = _ANCHOR.search(text) or _ANCHOR_LOOSE.search(text)
    return cn_to_int(m.group(1)) if m else None


def _norm(s: str) -> str:
    """Fold width, drop whitespace, combining marks and variation selectors.

    Word stores variation selectors inside words (`融為一︀體`), so two strings that read
    identically compare unequal without this.
    """
    s = unicodedata.normalize("NFKC", s or "")
    return "".join(
        c for c in s
        if not c.isspace()
        and unicodedata.category(c) not in ("Cf", "Mn")
        and not (0xE0100 <= ord(c) <= 0xE01EF)
    )


def _is_mark_cell(cell) -> list[str] | None:
    """The 段號 column: one numeral per paragraph, running 1, 2, 3, … with no gaps.

    Requiring the sequence to be complete and ascending is what separates it from a
    column of scores or a numbered list of questions.
    """
    marks = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    if len(marks) < _MIN_MARKS:
        return None
    if not all(len(m) <= 2 and m[0] in _NUMERAL_CHARS for m in marks):
        return None
    values = [cn_to_int(m) for m in marks]
    if None in values or values != list(range(1, len(values) + 1)):
        return None
    return marks


def read_numbered_body(docx_path: Path) -> tuple[list[str], list[str]] | None:
    """(段號 marks, 課文 cell paragraphs) from the body table, or None.

    Picks the row with a 段號 column and, within it, the longest DISTINCT cell as the
    text. The dedup is load-bearing: python-docx reports a horizontally merged cell once
    per grid column it spans, so `row.cells` yields the body text two or three times and
    "the longest cell" would otherwise be satisfied by a duplicate of the 段號 column in
    the lessons where that one is merged instead.
    """
    doc = docx.Document(str(docx_path))
    best: tuple[list[str], list[str]] | None = None
    for table in doc.tables:
        for row in table.rows:
            cells = list(row.cells)
            for i, cell in enumerate(cells):
                marks = _is_mark_cell(cell)
                if marks is None:
                    continue
                seen, candidates = set(), []
                for j, other in enumerate(cells):
                    if j == i:
                        continue
                    key = _norm(other.text)
                    if not key or key in seen:
                        continue
                    seen.add(key)
                    candidates.append(other)
                if not candidates:
                    continue
                text_cell = max(candidates, key=lambda c: len(_norm(c.text)))
                paras = [p.text.strip() for p in text_cell.paragraphs if p.text.strip()]
                if len(_norm("".join(paras))) < _MIN_CELL_CHARS:
                    continue
                if best is None or len(_norm("".join(paras))) > len(_norm("".join(best[1]))):
                    best = (marks, paras)
    return best


def _is_heading(text: str) -> bool:
    return len(_norm(text)) <= _HEADING_MAX and not text.rstrip().endswith(tuple(_SENTENCE_END))


def align_to_numbering(marks: list[str], cell_paras: list[str]) -> list[str] | None:
    """The cell's paragraphs, transformed until there are exactly as many as the author
    numbered — or None when no transformation gets there.

    Two transformations, tried in this order, each attempted only when the counts do not
    already agree:

      1. drop sub-headings — 說明文 prints 「無形的殺手」 above a run of paragraphs and
         does not number it. Dropping a short unpunctuated line cannot corrupt a
         paragraph's text, so this is the safe one.
      2. merge continuations — Word splits a paragraph at a manual line break, leaving a
         tail that starts mid-sentence (《感情小日記1》 lost 「有多一些互動，我就會覺得…」).
         This one CHANGES text, so it goes last.

    WHY THE COUNT IS THE ACCEPTANCE TEST, AND WHY THAT IS NOT ENOUGH
    ----------------------------------------------------------------
    Matching the author's own count is the only ground truth available here, and it is
    what makes each transformation confirmed rather than assumed. But it is necessary,
    not sufficient: two different transformations can reach the same count and the wrong
    one still "passes". That happened while this was being written — merging with 「……」
    absent from the sentence-end set hit the right count on 《感情小日記2》 by merging two
    genuinely separate paragraphs, and produced the wrong passage.

    So the transformations are ordered least-destructive-first, applied only when needed,
    and the result is still checked against the first edition's text where one exists.
    Measured on the 36 lessons 靖杭's golden set can judge: 29 → 31 correct, withheld
    4 → 2, and the 3 remaining errors are anchor-level (L0030 / L0072 / L0110) which no
    alignment can fix.
    """
    if 0 <= len(cell_paras) - len(marks) <= MAX_UNNUMBERED_TAIL:
        # One or two unnumbered tail lines; indexing lands regardless. Transforming here
        # would be solving a problem that is not present.
        return cell_paras

    dropped = [p for p in cell_paras if not _is_heading(p)]
    if len(dropped) == len(marks):
        return dropped

    merged: list[str] = []
    for para in dropped:
        if merged and not merged[-1].rstrip().endswith(tuple(_SENTENCE_END)):
            merged[-1] = merged[-1] + para
        else:
            merged.append(para)
    if len(merged) == len(marks):
        return merged
    return None


def load_legacy() -> list[str]:
    """First-edition passages, TEXT only.

    The paragraph numbers are deliberately discarded. Keeping them is what let the old
    corroboration compare 「第三段」 in one edition's printing against an index into the
    other edition's derived paragraph list and call the two agreeing a confirmation.
    """
    import yaml

    if not LEGACY_TABLE.exists():
        return []
    doc = yaml.safe_load(LEGACY_TABLE.read_text(encoding="utf-8")) or {}
    out = []
    for entry in (doc.get("passages") or {}).values():
        passage = (entry or {}).get("passage")
        if passage:
            out.append(passage.strip())
    return out


def corroborate(passage: str, body: list[str], legacy: list[str]) -> bool | None:
    """Does the first edition agree with the passage we are about to write?

    True  — a first-edition passage belongs to this lesson AND is this passage.
    False — one belongs to this lesson and is a DIFFERENT paragraph.
    None  — this lesson has no first-edition counterpart (its text was rewritten, or it
            is new). Not a pass: the caller must not treat None as agreement.

    "Belongs to this lesson" is decided by content, not by lesson code — the codes are
    what made the old lookup serve another lesson's paragraph. Ambiguous matches (a
    passage occurring in two lessons) are discarded rather than guessed.
    """
    joined = _norm("".join(body))
    mine = _norm(passage)
    hits = [p for p in legacy if _norm(p)[:60] and _norm(p)[:60] in joined]
    if not hits:
        return None
    if len({_norm(p) for p in hits}) > 1:
        return None
    return _norm(hits[0]) == mine


def _is_body_span(passage: str, body: list[str]) -> bool:
    """Is `passage` one stored paragraph, or a run of consecutive stored paragraphs?"""
    target = _norm(passage)
    normed = [_norm(p) for p in body]
    for i in range(len(normed)):
        acc = ""
        for j in range(i, len(normed)):
            acc += normed[j]
            if acc == target:
                return True
            if len(acc) > len(target):
                break
    return False


def extract(docx_path: Path, body: list[str] | None = None) -> dict:
    """`body` should be the paragraphs as STORED, not a fresh re-extraction.

    The stored body is what the student reads and what the TTS sentence table is built
    from, so a passage that is not one of its paragraphs cannot be served whole even if
    the numbering says it is right.
    """
    docx_path = Path(docx_path)
    paras = _paragraphs(docx_path)
    if body is None:
        body = (extract_body(docx_path) or {}).get("paragraphs") or []

    out: dict = {
        "anchor": None,
        "verdict": "empty",
        "passage": None,
        "corroborated_by_first_edition": None,
        "needs_human_review": False,
    }

    # Set before every early return below: the fluency target belongs to the 念順順
    # section, not to the paragraph, so a lesson keeps its target even when the anchor
    # is withheld (#2722). The 12 文言文 lessons time in seconds and mostly have no
    # anchor at all.
    benchmark = find_reading_benchmark(paras)
    if benchmark:
        out["reading_benchmark"] = benchmark

    anchor = find_anchor(paras)
    out["anchor"] = anchor
    if anchor is None:
        out["verdict"] = "no_anchor"
        return out

    numbered = read_numbered_body(docx_path)
    if numbered is None:
        # No 段號 column found. The anchor names a printed paragraph number and there is
        # nothing here that says which paragraph carries it; indexing the derived body
        # instead is exactly the substitution that produced #2720.
        out["verdict"] = "no_printed_numbering"
        return out

    marks, cell_paras = numbered
    out["printed_marks"] = len(marks)
    out["cell_paragraphs"] = len(cell_paras)

    out["unnumbered_tail"] = len(cell_paras) - len(marks)
    aligned = align_to_numbering(marks, cell_paras)
    if aligned is None:
        # No transformation reaches the author's count, so "the Nth paragraph" is not a
        # well-defined thing here. Withheld rather than guessed at.
        out["verdict"] = "numbering_disagrees"
        out["needs_human_review"] = True
        return out
    out["aligned_paragraphs"] = len(aligned)
    cell_paras = aligned

    if anchor > len(cell_paras):
        out["verdict"] = "anchor_out_of_range"
        out["needs_human_review"] = True
        return out

    passage = cell_paras[anchor - 1].strip()
    out["passage"] = passage
    out["start_text"] = passage[:24]
    out["paragraphs_used"] = 1

    if not (MIN_CHARS <= len(_norm(passage)) <= MAX_CHARS):
        out["verdict"] = "implausible_length"
        out["length"] = len(_norm(passage))
        out["needs_human_review"] = True
        return out

    if not body:
        out["verdict"] = "no_body"
        return out

    if not _is_body_span(passage, body):
        # The passage has to be text the student is actually shown. One printed paragraph
        # can span two of `body.yml`'s (Word split it at a manual line break), which is
        # why a RUN of consecutive stored paragraphs counts — serving only the first
        # would cut 《感情小日記1》 off mid-sentence. Anything that is not such a run
        # means the body extractor and the numbering disagree about the text itself, and
        # that is withheld rather than trimmed to fit.
        out["verdict"] = "not_a_stored_paragraph"
        out["needs_human_review"] = True
        return out

    agreed = corroborate(passage, body, load_legacy())
    out["corroborated_by_first_edition"] = agreed
    if agreed is False:
        # 二修教材為主 (靖杭, 2026-08-18). The second edition's own instruction about the
        # second edition's worksheet outranks a passage read off the first edition's
        # printing, so a disagreement no longer withholds — it is written and FLAGGED.
        #
        # The cost is measured, not assumed: on the 36 judgeable lessons this writes 3
        # passages (L0030 / L0072 / L0110) that the first edition says are the wrong
        # paragraph, and in those three the first edition's text still appears verbatim
        # in the second edition's body — so the text did not change and the disagreement
        # is more likely our misread than a re-marking. They are listed in
        # `docs/curriculum/key-reading-disagrees-first-edition.md` for a human to settle.
        out["verdict"] = "disagrees_with_first_edition"
        out["needs_human_review"] = True
        return out

    out["verdict"] = "confirmed" if agreed else "ok"
    return out


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--source", default="/tmp/docx-src")
    ap.add_argument("--out", default="/tmp/key_reading.json")
    a = ap.parse_args()

    results, counts = {}, {}
    for docx_file in sorted(Path(a.source).glob("*.docx")):
        r = extract(docx_file)
        results[docx_file.stem] = r
        counts[r["verdict"]] = counts.get(r["verdict"], 0) + 1

    Path(a.out).write_text(
        json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(f"  {len(results)} 課")
    for v, n in sorted(counts.items(), key=lambda kv: -kv[1]):
        print(f"    {v:30s} {n}")
    ok = [r for r in results.values() if r["verdict"] in ("ok", "confirmed")]
    if ok:
        conf = sum(1 for r in ok if r["verdict"] == "confirmed")
        lengths = sorted(len(r["passage"]) for r in ok)
        print(f"\n  可寫入 {len(ok)}（其中 {conf} 課與一修的段落文字相符）")
        print(f"  段落長度 中位 {lengths[len(lengths) // 2]} 字")
    return 0


if __name__ == "__main__":
    sys.exit(main())
