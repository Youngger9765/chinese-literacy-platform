#!/usr/bin/env python3
"""
Parse 世杰老師自學教材 Word files into structured JSON.

Usage:
    python3 scripts/parse_materials.py                          # parse all
    python3 scripts/parse_materials.py --file path/to/G4-1.docx # parse one
    python3 scripts/parse_materials.py --grade 4                # parse one grade

Output: data/parsed/{grade}/{G4-1-L1.json, ...}
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from pathlib import Path
from typing import Optional

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx required. Install with: pip install python-docx")
    sys.exit(1)

try:
    import openpyxl
except ImportError:
    openpyxl = None

RAW_BASE = Path("data/raw-materials/三民國中自學教材/1.分享三民國中－自學教材")
WORD_DIR = RAW_BASE / "1-1.國中自學教材L1-60學習單0916" / "1-2教用版WORD"
OVERVIEW_FILE = RAW_BASE / "3.L1-60教材總覽及閱聚光燈策略層次0916.xlsx"
OUTPUT_DIR = Path("data/parsed")


def load_metadata() -> dict:
    """Load lesson metadata from overview spreadsheet."""
    if not openpyxl or not OVERVIEW_FILE.exists():
        return {}

    wb = openpyxl.load_workbook(OVERVIEW_FILE, read_only=True)
    ws = wb[wb.sheetnames[0]]

    meta = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not row[0]:
            continue
        lesson_num = int(row[0])
        meta[lesson_num] = {
            "grade_code": str(row[1] or ""),
            "text_type": str(row[2] or ""),   # 單/多
            "title": str(row[3] or ""),
            "genre": str(row[4] or ""),        # 記敘文/說明文
            "reading_strategy": str(row[5] or ""),
            "grade": int(row[6]) if row[6] else None,
        }
    return meta


def extract_story_from_table(doc: Document) -> tuple[str, list[str]]:
    """Extract story text from Table 0 (standard format).
    For multi-text lessons, concatenates stories from multiple tables."""
    if not doc.tables:
        return "", []

    all_paragraphs = []

    for t in doc.tables:
        # Story tables: 1 row, 1-2 columns, cell(0,0) has >50 chars of text
        if len(t.rows) != 1 or len(t.columns) > 2:
            continue
        cell_text = t.cell(0, 0).text.strip()
        cell_start = cell_text[:20]
        if len(cell_text) < 50 or "計時" in cell_start or "影片" in cell_start:
            continue

        for chunk in re.split(r'\n\s*\n|\n(?=　)', cell_text):
            chunk = chunk.strip()
            if chunk and not re.match(r'^\d+$', chunk):
                all_paragraphs.append(chunk)

    full_text = "\n".join(all_paragraphs)
    return full_text, all_paragraphs


def extract_story_from_paragraphs(doc: Document) -> tuple[str, list[str]]:
    """Extract story from paragraphs (classical text format)."""
    paragraphs = []
    in_story = False
    story_ended = False

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        # Look for lesson title to start
        if re.match(r'^第\d+課', text):
            in_story = True
            continue

        if in_story and not story_ended:
            # Stop at reading exercise markers
            if any(kw in text for kw in ['計時器', '朗讀', '自我練習', '自我檢核', '裁判']):
                story_ended = True
                continue
            # Skip author attribution lines
            if re.match(r'^(清|明|宋|唐|漢)·', text):
                continue
            if text.startswith('注釋') or text.startswith('語譯'):
                continue
            if len(text) > 10:
                paragraphs.append(text)

    full_text = "\n".join(paragraphs)
    return full_text, paragraphs


def extract_vocabulary(doc: Document) -> list[dict]:
    """Extract vocabulary words with definitions."""
    vocab = []
    in_vocab = False

    for p in doc.paragraphs:
        text = p.text.strip()

        # Detect vocab section start
        if '填入本課語詞' in text or '空格內填入' in text:
            in_vocab = True
            continue

        if in_vocab:
            # Stop at next section
            if '選出最適合的答案' in text or '根據文章內容' in text:
                break
            if '將正確的詞語代號' in text or '填入空格' in text:
                break

            # Parse "(1) 詞語 ：解釋" pattern
            m = re.match(r'^\((\d+)\)\s*(.+?)\s*[：:]\s*(.+)$', text)
            if m:
                vocab.append({
                    "word": m.group(2).strip(),
                    "definition": m.group(3).strip(),
                })
            elif text.startswith('*') or text.startswith('＊'):
                if vocab:
                    vocab[-1]["note"] = text.lstrip('*＊').strip()

    # Fallback: extract 多義字 from tables (classical texts like G8)
    if not vocab:
        vocab = _extract_vocab_from_tables(doc)

    # Fallback: extract inline definitions from paragraphs (classical texts)
    if not vocab:
        vocab = _extract_inline_definitions(doc)

    return vocab


def _extract_vocab_from_tables(doc: Document) -> list[dict]:
    """Extract vocabulary from 多義字 tables (used in classical text lessons)."""
    vocab = []
    for t in doc.tables:
        cell00 = t.cell(0, 0).text.strip()
        if '多義字' not in cell00 and '字義' not in cell00:
            continue
        # Multi-meaning char tables: col 0 = char, other cols = meanings
        for ri in range(t.rows.__len__()):
            try:
                word = t.cell(ri, 0).text.strip()
                if word in ('多義字', '字義', '文章結構', '') or len(word) > 4:
                    continue
                meanings = []
                for ci in range(1, len(t.columns)):
                    m = t.cell(ri, ci).text.strip()
                    if m and m not in ('', word):
                        meanings.append(m)
                if meanings:
                    vocab.append({
                        "word": word,
                        "definition": "；".join(meanings),
                        "type": "多義字",
                    })
            except Exception:
                continue
    return vocab


def _extract_inline_definitions(doc: Document) -> list[dict]:
    """Extract inline definitions like '雨，當動詞，表示落下的意思' from paragraphs."""
    vocab = []
    for p in doc.paragraphs:
        text = p.text.strip()
        # Pattern: "word：definition" or "word，meaning的意思"
        m = re.match(r'^(\S{1,4})[：:，]\s*(?:有「|當|表示|意思是)(.+?)(?:的意思|。|$)', text)
        if m and len(m.group(1)) <= 4:
            vocab.append({
                "word": m.group(1),
                "definition": m.group(2).strip().rstrip('。'),
                "type": "inline",
            })
        # Pattern: "word：有「meaning」的意思"
        m2 = re.match(r'^(\S{1,4})[：:]有「(.+?)」的意思', text)
        if m2:
            vocab.append({
                "word": m2.group(1),
                "definition": m2.group(2),
                "type": "inline",
            })
    return vocab


def extract_fill_in_blank(doc: Document) -> list[dict]:
    """Extract fill-in-the-blank exercises."""
    exercises = []
    in_section = False
    options_text = ""

    for p in doc.paragraphs:
        text = p.text.strip()

        if '將正確的詞語代號' in text:
            in_section = True
            continue

        if in_section:
            # Stop at next section
            if '選出最適合的答案' in text or '根據文章內容' in text:
                break

            # Options line (A疑難雜症 B龍爭虎鬥 ...)
            if re.match(r'^[A-GＡ-Ｇ]', text):
                options_text += " " + text
                continue

            # Exercise lines: (1)句子 (  D  )
            m = re.match(r'^\((\d+)\)(.+)', text)
            if m:
                sentence = m.group(2).strip()
                # Extract answer from parentheses
                ans_match = re.search(r'\(\s*([A-GＡ-Ｇ])\s*\)', sentence)
                answer = ans_match.group(1) if ans_match else None
                exercises.append({
                    "sentence": re.sub(r'\(\s*[A-GＡ-Ｇ]?\s*\)', '(　　)', sentence),
                    "answer": answer,
                })

    return exercises


def extract_multiple_choice(doc: Document) -> list[dict]:
    """Extract multiple choice comprehension questions."""
    questions = []
    current_q = None
    in_mc_section = False

    def _save_q():
        nonlocal current_q
        if current_q and current_q["options"]:
            questions.append(current_q)
        current_q = None

    def _new_q(q_text):
        nonlocal current_q
        _save_q()
        current_q = {
            "question": q_text,
            "options": [],
            "answer": None,
            "explanation": None,
        }

    def _add_option(raw):
        if current_q is None:
            return
        # Strip option prefix: (A) or A. or Ａ．
        clean = re.sub(
            r'^(?:[\(（][A-DＡ-Ｄ][\)）]|[A-DＡ-Ｄ][\.\．、]?)\s*', '', raw
        ).strip()
        if not clean:
            return
        # Extract trailing teacher explanation （...3+ chars...）
        exp = re.search(r'[（(]([^A-DＡ-Ｄ].{3,}?)[）)]\s*$', clean)
        if exp:
            current_q["answer"] = chr(ord('A') + len(current_q["options"]))
            current_q["explanation"] = exp.group(1)
            clean = clean[:exp.start()].strip()
        if clean:
            current_q["options"].append(clean)

    # Build line list — split multi-line paragraphs so embedded options
    # (e.g. "question\n(A)option" in one paragraph) are processed individually
    lines = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        raw = p.text.strip()
        if not raw:
            continue
        for line in raw.split('\n'):
            line = line.strip()
            if line:
                lines.append((line, style))

    for text, style in lines:
        # Detect MC section start (broader matching)
        if ('選出' in text and ('答案' in text or '適合' in text or '正確' in text)) \
                or ('根據' in text and '選出' in text) \
                or text == '閱讀理解測驗':
            in_mc_section = True
            continue

        # Stop at non-MC sections (word search, video links)
        if in_mc_section and ('找一找' in text or '影片連結' in text or '圈出' in text):
            break

        # --- Question patterns ---
        # Classical: （  ）N.question
        classical_q = re.match(r'^[（(][　 ]*[）)]\s*(\d+)[\.\．]\s*(.+)', text)
        # Multi-text: （　　）(N)question
        multi_text_q = re.match(r'^[（(][　 ]*[）)]\s*[\(（](\d+)[\)）]\s*(.+)', text)
        # Reversed: (N)（　　）question
        reversed_q = re.match(r'^[\(（](\d+)[\)）]\s*[（(][　 ]*[）)]\s*(.+)', text)

        # --- Option patterns ---
        is_paren_opt = bool(re.match(r'^[\(（][A-DＡ-Ｄ][\)）]', text))
        is_dot_opt = bool(re.match(r'^[A-DＡ-Ｄ][\.\．、]', text))
        # Handle missing dot: "D未解之謎" (letter + CJK, no dot)
        is_bare_opt = bool(re.match(r'^[A-DＡ-Ｄ](?=[\u4e00-\u9fff])', text))
        is_option = is_paren_opt or is_dot_opt or is_bare_opt

        # Multi-option line: "A.xxx  B.xxx  C.xxx  D.xxx"
        is_multi_opt = (is_dot_opt or is_bare_opt) and bool(
            re.search(r'[\s　][B-DＢ-Ｄ][\.\．、]', text)
        )

        # Question keyword detection (exclude option lines)
        is_question = (
            ('?' in text or '？' in text or '為何' in text or '請問' in text
             or '哪' in text or '何者' in text or '下列' in text)
            and not is_option
            and len(text) > 10
        )

        # --- Process (order matters: specific patterns first) ---
        if multi_text_q:
            _new_q(multi_text_q.group(2))
        elif reversed_q:
            _new_q(reversed_q.group(2))
        elif classical_q and in_mc_section:
            _new_q(classical_q.group(2))
        elif is_question and ('?' in text or '？' in text) and style == "List Paragraph":
            _new_q(text)
        elif is_question and in_mc_section:
            _new_q(text)
        elif is_multi_opt and current_q is not None:
            # Split "A.xxx  B.xxx" into individual options
            parts = re.split(r'[\s　]+(?=[A-DＡ-Ｄ][\.\．、])', text)
            for part in parts:
                _add_option(part.strip())
        elif is_option and current_q is not None:
            _add_option(text)

    _save_q()
    return questions


def extract_reading_benchmark(doc: Document) -> Optional[dict]:
    """Extract reading speed benchmarks from self-check table."""
    for t in doc.tables:
        cell00 = t.cell(0, 0).text.strip()
        # Look for the rubric table (□＜190字 pattern)
        if '□' in cell00 or '字' in cell00:
            if t.rows and t.columns and len(t.rows) == 2 and len(t.columns) == 3:
                try:
                    levels = []
                    for ci in range(3):
                        threshold = t.cell(0, ci).text.strip()
                        feedback = t.cell(1, ci).text.strip()
                        levels.append({"threshold": threshold, "feedback": feedback})
                    return {"levels": levels}
                except Exception:
                    pass
    return None


def parse_filename(filename: str) -> dict:
    """Extract grade and lesson info from filename."""
    # G4-1-L1贏得喝采的輸家.docx
    m = re.match(r'G(\d+)-(\d+)-L(\d+)(.+?)\.docx', filename)
    if m:
        return {
            "grade": int(m.group(1)),
            "order": int(m.group(2)),
            "lesson_number": int(m.group(3)),
            "title": m.group(4),
        }
    # Multi-text: G9-11.12.13-L58~60多文本-...
    m = re.match(r'G(\d+)-([\d.]+)-L(\d+)~?(\d+)?(.+?)\.docx', filename)
    if m:
        return {
            "grade": int(m.group(1)),
            "order": m.group(2),
            "lesson_number": int(m.group(3)),
            "lesson_end": int(m.group(4)) if m.group(4) else None,
            "title": m.group(5),
            "is_multi_text": True,
        }
    return {"title": filename}


def parse_docx(filepath: Path, metadata: dict) -> dict:
    """Parse a single docx file into structured data."""
    doc = Document(str(filepath))
    file_info = parse_filename(filepath.name)

    # Try standard story extraction (Table 0)
    story_text, paragraphs = extract_story_from_table(doc)

    # Fallback: extract from paragraphs (classical texts)
    if not story_text:
        story_text, paragraphs = extract_story_from_paragraphs(doc)

    # Extract components
    vocabulary = extract_vocabulary(doc)
    fill_in_blank = extract_fill_in_blank(doc)
    multiple_choice = extract_multiple_choice(doc)
    reading_benchmark = extract_reading_benchmark(doc)

    # Get metadata from overview
    lesson_num = file_info.get("lesson_number")
    meta = metadata.get(lesson_num, {})

    result = {
        "lesson_number": lesson_num,
        "grade": file_info.get("grade"),
        "grade_code": meta.get("grade_code", f"G{file_info.get('grade', '?')}-{file_info.get('order', '?')}"),
        "title": file_info.get("title", ""),
        "genre": meta.get("genre", ""),
        "text_type": meta.get("text_type", "單"),
        "reading_strategy": meta.get("reading_strategy", ""),
        "story_text": story_text,
        "paragraphs": paragraphs,
        "paragraph_count": len(paragraphs),
        "char_count": len(story_text.replace("\n", "").replace(" ", "")),
        "vocabulary": vocabulary,
        "vocabulary_count": len(vocabulary),
        "fill_in_blank": fill_in_blank,
        "multiple_choice": multiple_choice,
        "multiple_choice_count": len(multiple_choice),
        "reading_benchmark": reading_benchmark,
        "source_file": filepath.name,
    }

    # Quality flags
    flags = []
    if not story_text:
        flags.append("NO_STORY_TEXT")
    if not vocabulary:
        flags.append("NO_VOCABULARY")
    if not multiple_choice:
        flags.append("NO_MC_QUESTIONS")
    if file_info.get("is_multi_text"):
        flags.append("MULTI_TEXT")
    result["flags"] = flags

    return result


def parse_all(grade_filter: Optional[int] = None):
    """Parse all Word files and output JSON."""
    metadata = load_metadata()

    grades = [4, 5, 6, 7, 8, 9]
    if grade_filter:
        grades = [grade_filter]

    total = 0
    success = 0
    flagged = 0

    for grade in grades:
        grade_dir = WORD_DIR / f"{grade}年級"
        if not grade_dir.exists():
            print(f"  SKIP: {grade_dir} not found")
            continue

        out_dir = OUTPUT_DIR / f"grade-{grade}"
        out_dir.mkdir(parents=True, exist_ok=True)

        files = sorted(grade_dir.glob("*.docx"))
        print(f"\nGrade {grade}: {len(files)} files")

        for fpath in files:
            total += 1
            try:
                result = parse_docx(fpath, metadata)
                out_file = out_dir / f"{fpath.stem}.json"
                with open(out_file, "w", encoding="utf-8") as f:
                    json.dump(result, f, ensure_ascii=False, indent=2)

                flags = result.get("flags", [])
                if flags:
                    flagged += 1
                    print(f"  ⚠ {fpath.name}: {', '.join(flags)}")
                else:
                    print(f"  ✓ {fpath.name} → {result['char_count']} chars, "
                          f"{result['vocabulary_count']} vocab, "
                          f"{result['multiple_choice_count']} MC")
                success += 1

            except Exception as e:
                print(f"  ✗ {fpath.name}: {e}")

    print(f"\n{'='*50}")
    print(f"Total: {total} | Success: {success} | Flagged: {flagged} | Failed: {total - success}")


def parse_single(filepath: str):
    """Parse a single file and print JSON."""
    metadata = load_metadata()
    result = parse_docx(Path(filepath), metadata)
    print(json.dumps(result, ensure_ascii=False, indent=2))


def main():
    parser = argparse.ArgumentParser(description="Parse teaching materials to JSON")
    parser.add_argument("--file", help="Parse a single docx file")
    parser.add_argument("--grade", type=int, help="Parse one grade (4-9)")
    args = parser.parse_args()

    if args.file:
        parse_single(args.file)
    else:
        parse_all(args.grade)


if __name__ == "__main__":
    main()
